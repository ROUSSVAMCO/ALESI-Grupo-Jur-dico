
import os
import smtplib
import subprocess
import tempfile
import shutil
import re
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from email.message import EmailMessage
from urllib.parse import quote_plus
from io import BytesIO
from openpyxl import Workbook

from flask import Flask, render_template_string, request, redirect, flash, session, send_from_directory, Response, send_file
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import text
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

try:
    import cloudinary
    import cloudinary.uploader
except Exception:
    cloudinary = None


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "ALESI_CAMBIAR_SECRET_KEY")

DB = os.environ.get("DATABASE_URL", "sqlite:///alesi.db")
if DB.startswith("postgres://"):
    DB = DB.replace("postgres://", "postgresql://", 1)

app.config["SQLALCHEMY_DATABASE_URI"] = DB
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

os.makedirs("uploads", exist_ok=True)


def local_now():
    return datetime.now(ZoneInfo("America/Tijuana")).replace(tzinfo=None)


def local_date_iso():
    return local_now().strftime("%Y-%m-%d")


def local_stamp():
    return local_now().strftime("%Y%m%d%H%M%S")


MATERIAS = [
    "Familiar", "Laboral", "Civil", "Mercantil", "Penal", "Administrativo",
    "Fiscal", "Amparo", "Agrario", "Corporativo", "Notarial", "Municipal", "Otro"
]

TIPOS_INICIALES = [
    ("Familiar", "Divorcio voluntario"), ("Familiar", "Divorcio incausado"),
    ("Familiar", "Pensión alimenticia"), ("Familiar", "Aumento de pensión"),
    ("Familiar", "Reducción de pensión"), ("Familiar", "Guarda y custodia"),
    ("Familiar", "Régimen de convivencia"), ("Familiar", "Patria potestad"),
    ("Familiar", "Violencia familiar"), ("Familiar", "Sucesión testamentaria"),
    ("Familiar", "Sucesión intestamentaria"), ("Familiar", "Juicio hereditario"),
    ("Laboral", "Despido injustificado"), ("Laboral", "Pago de prestaciones"),
    ("Laboral", "Horas extras"), ("Laboral", "Reinstalación"),
    ("Laboral", "Convenio laboral"), ("Laboral", "Accidente de trabajo"),
    ("Civil", "Cumplimiento de contrato"), ("Civil", "Incumplimiento de contrato"),
    ("Civil", "Daños y perjuicios"), ("Civil", "Cobro de pesos"),
    ("Civil", "Arrendamiento"), ("Civil", "Prescripción positiva"),
    ("Civil", "Otorgamiento y firma de escritura"),
    ("Mercantil", "Cobro de pagaré"), ("Mercantil", "Juicio ejecutivo mercantil"),
    ("Mercantil", "Juicio oral mercantil"), ("Mercantil", "Cobro de factura"),
    ("Penal", "Robo"), ("Penal", "Fraude"), ("Penal", "Abuso de confianza"),
    ("Penal", "Daño en propiedad ajena"), ("Penal", "Despojo"),
    ("Penal", "Amenazas"), ("Penal", "Lesiones"),
    ("Penal", "Incumplimiento de obligaciones alimentarias"),
    ("Administrativo", "Responsabilidad administrativa"),
    ("Administrativo", "Procedimiento administrativo"), ("Administrativo", "Revocación de multa"),
    ("Administrativo", "Clausura"), ("Administrativo", "Auditoría"),
    ("Fiscal", "Crédito fiscal"), ("Fiscal", "Multa fiscal"), ("Fiscal", "Recurso de revocación"),
    ("Fiscal", "Juicio contencioso administrativo"),
    ("Amparo", "Amparo indirecto"), ("Amparo", "Amparo directo"),
    ("Amparo", "Suspensión provisional"), ("Amparo", "Suspensión definitiva"),
    ("Agrario", "Conflicto parcelario"), ("Agrario", "Posesión agraria"),
    ("Corporativo", "Constitución de sociedad"), ("Corporativo", "Acta de asamblea"),
    ("Corporativo", "Poderes"), ("Corporativo", "Modificación estatutaria"),
    ("Notarial", "Notario"), ("Notarial", "Copias certificadas"),
    ("Municipal", "Permiso municipal"), ("Municipal", "Uso de suelo"),
    ("Otro", "Convenio extrajudicial"), ("Otro", "Carta poder"), ("Otro", "Regularización de predio")
]


class Usuario(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nombre = db.Column(db.String(180), nullable=False)
    usuario = db.Column(db.String(80), unique=True, nullable=False)
    correo = db.Column(db.String(180))
    password_hash = db.Column(db.String(255), nullable=False)
    rol = db.Column(db.String(80), default="Usuario")
    area = db.Column(db.String(80), default="")
    foto_url = db.Column(db.Text)
    activo = db.Column(db.Boolean, default=True)
    creado_en = db.Column(db.DateTime, default=datetime.now)


class TipoAsunto(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    materia = db.Column(db.String(80), nullable=False)
    nombre = db.Column(db.String(180), nullable=False)
    activo = db.Column(db.Boolean, default=True)


class Cliente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nombre = db.Column(db.String(200), nullable=False)
    telefono = db.Column(db.String(80))
    correo = db.Column(db.String(180))
    rfc = db.Column(db.String(80))
    curp = db.Column(db.String(80))
    materia = db.Column(db.String(80))
    tipo_asunto = db.Column(db.String(180))
    cobro_total = db.Column(db.Float, default=0)
    moneda_cobro = db.Column(db.String(10), default="MXN")
    direccion = db.Column(db.Text)
    observaciones = db.Column(db.Text)
    activo = db.Column(db.Boolean, default=True)
    creado_en = db.Column(db.DateTime, default=datetime.now)


class Expediente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    propietario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"), nullable=False)
    cliente_id = db.Column(db.Integer, db.ForeignKey("cliente.id"))
    numero = db.Column(db.String(120), nullable=False)
    materia = db.Column(db.String(80))
    tipo_asunto = db.Column(db.String(180))
    cobro_total = db.Column(db.Float, default=0)
    moneda_cobro = db.Column(db.String(10), default="MXN")
    abono_pactado = db.Column(db.Float, default=0)
    periodicidad_abono = db.Column(db.String(80), default="")
    autoridad = db.Column(db.String(200))
    actor = db.Column(db.String(200))
    demandado = db.Column(db.String(200))
    estado = db.Column(db.String(80), default="En trámite")
    prioridad = db.Column(db.String(40), default="Media")
    responsable = db.Column(db.String(180))
    fecha_inicio = db.Column(db.String(20))
    observaciones = db.Column(db.Text)
    creado_en = db.Column(db.DateTime, default=datetime.now)

    propietario = db.relationship("Usuario")
    cliente = db.relationship("Cliente")


class Compartido(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    permiso = db.Column(db.String(40), default="Lectura")
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class Movimiento(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    titulo = db.Column(db.String(220), nullable=False)
    fecha = db.Column(db.String(20))
    estatus = db.Column(db.String(80), default="Pendiente")
    proxima_accion = db.Column(db.String(250))
    fecha_limite = db.Column(db.String(20))
    hora_limite = db.Column(db.String(10), default="09:00")
    observaciones = db.Column(db.Text)
    archivo_url = db.Column(db.Text)
    notificar = db.Column(db.Boolean, default=True)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class Audiencia(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    titulo = db.Column(db.String(220), nullable=False)
    fecha = db.Column(db.String(20), nullable=False)
    hora = db.Column(db.String(10), nullable=False)
    autoridad = db.Column(db.String(220))
    sala = db.Column(db.String(120))
    modalidad = db.Column(db.String(80), default="Presencial")
    enlace = db.Column(db.Text)
    observaciones = db.Column(db.Text)
    notificar = db.Column(db.Boolean, default=True)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class Cobranza(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    folio = db.Column(db.String(40), unique=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey("cliente.id"))
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    concepto = db.Column(db.String(120))
    descripcion = db.Column(db.Text)
    monto_total = db.Column(db.Float, default=0)
    abono = db.Column(db.Float, default=0)
    saldo = db.Column(db.Float, default=0)
    moneda = db.Column(db.String(10), default="MXN")
    forma_pago = db.Column(db.String(50), default="Efectivo")
    fecha_pago = db.Column(db.String(20))
    proximo_pago = db.Column(db.String(20))
    estatus = db.Column(db.String(40), default="Pendiente")
    comprobante_url = db.Column(db.Text)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    cliente = db.relationship("Cliente")
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class NotificacionEnviada(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    tipo = db.Column(db.String(40))
    referencia_id = db.Column(db.Integer)
    anticipacion = db.Column(db.String(20))
    usuario_id = db.Column(db.Integer)
    enviado_en = db.Column(db.DateTime, default=datetime.now)


class Bitacora(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    accion = db.Column(db.String(220), nullable=False)
    detalle = db.Column(db.Text)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class MensajeExpediente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    mensaje = db.Column(db.Text, nullable=False)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class NotificacionInterna(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    titulo = db.Column(db.String(220))
    mensaje = db.Column(db.Text)
    enlace = db.Column(db.String(250))
    leida = db.Column(db.Boolean, default=False)
    creador_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    grupo = db.Column(db.String(80))
    creado_en = db.Column(db.DateTime, default=datetime.now)
    usuario = db.relationship("Usuario", foreign_keys=[usuario_id])
    creador = db.relationship("Usuario", foreign_keys=[creador_id])


class AvisoExpediente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    titulo = db.Column(db.String(220))
    mensaje = db.Column(db.Text)
    fecha_aviso = db.Column(db.String(20))
    creado_en = db.Column(db.DateTime, default=datetime.now)
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class ContratoCliente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey("cliente.id"))
    expediente_id = db.Column(db.Integer, db.ForeignKey("expediente.id"))
    usuario_id = db.Column(db.Integer, db.ForeignKey("usuario.id"))
    cliente_nombre = db.Column(db.String(200))
    tipo_asunto = db.Column(db.String(180))
    monto_total = db.Column(db.Float, default=0)
    monto_letra = db.Column(db.Text)
    forma_pago = db.Column(db.Text)
    fecha_contrato = db.Column(db.String(40))
    personal_atendio = db.Column(db.String(200))
    telefono_cliente = db.Column(db.String(80))
    correo_cliente = db.Column(db.String(180))
    creado_en = db.Column(db.DateTime, default=datetime.now)
    cliente = db.relationship("Cliente")
    expediente = db.relationship("Expediente")
    usuario = db.relationship("Usuario")


class CitaProspecto(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey("cliente.id"))
    nombre = db.Column(db.String(200), nullable=False)
    telefono = db.Column(db.String(80))
    correo = db.Column(db.String(180))
    materia = db.Column(db.String(80))
    tipo_asunto = db.Column(db.String(180))
    fecha = db.Column(db.String(20))
    hora = db.Column(db.String(10))
    documentacion_solicitada = db.Column(db.Text)
    documentacion_original = db.Column(db.Text)
    estatus = db.Column(db.String(80), default="Agendada")
    observaciones = db.Column(db.Text)
    creado_en = db.Column(db.DateTime, default=datetime.now)
    cliente = db.relationship("Cliente")


class Apariencia(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    color_principal = db.Column(db.String(20), default="#111827")
    color_secundario = db.Column(db.String(20), default="#14B8A6")
    color_dorado = db.Column(db.String(20), default="#B38B2E")
    fuente = db.Column(db.String(80), default="Arial")
    fondo = db.Column(db.String(20), default="#F8FAFC")
    logo_url = db.Column(db.Text)


def ensure_schema_updates():
    dialect = db.engine.dialect.name

    def add_column(table, column, definition):
        try:
            if dialect == "postgresql":
                db.session.execute(text(f"ALTER TABLE {table} ADD COLUMN IF NOT EXISTS {column} {definition}"))
            elif dialect == "sqlite":
                cols = [row[1] for row in db.session.execute(text(f"PRAGMA table_info({table})")).fetchall()]
                if column not in cols:
                    db.session.execute(text(f"ALTER TABLE {table} ADD COLUMN {column} {definition}"))
        except Exception as exc:
            print("Schema update skipped:", table, column, exc)

    add_column("usuario", "area", "VARCHAR(80)")
    add_column("usuario", "foto_url", "TEXT")
    add_column("cliente", "rfc", "VARCHAR(80)")
    add_column("cliente", "curp", "VARCHAR(80)")
    add_column("cliente", "materia", "VARCHAR(80)")
    add_column("cliente", "tipo_asunto", "VARCHAR(180)")
    add_column("cliente", "cobro_total", "FLOAT DEFAULT 0")
    add_column("cliente", "moneda_cobro", "VARCHAR(10) DEFAULT 'MXN'")
    add_column("cliente", "activo", "BOOLEAN DEFAULT TRUE")
    add_column("expediente", "materia", "VARCHAR(80)")
    add_column("expediente", "tipo_asunto", "VARCHAR(180)")
    add_column("expediente", "cobro_total", "FLOAT DEFAULT 0")
    add_column("expediente", "moneda_cobro", "VARCHAR(10) DEFAULT 'MXN'")
    add_column("expediente", "abono_pactado", "FLOAT DEFAULT 0")
    add_column("expediente", "periodicidad_abono", "VARCHAR(80) DEFAULT ''")
    add_column("cobranza", "folio", "VARCHAR(40)")

    # Avisos internos globales
    add_column("notificacion_interna", "creador_id", "INTEGER")
    add_column("notificacion_interna", "grupo", "VARCHAR(80)")

    # Apariencia / personalización del sistema
    add_column("apariencia", "color_principal", "VARCHAR(20) DEFAULT '#111827'")
    add_column("apariencia", "color_secundario", "VARCHAR(20) DEFAULT '#14B8A6'")
    add_column("apariencia", "color_dorado", "VARCHAR(20) DEFAULT '#B38B2E'")
    add_column("apariencia", "fuente", "VARCHAR(80) DEFAULT 'Arial'")
    add_column("apariencia", "fondo", "VARCHAR(20) DEFAULT '#F8FAFC'")
    add_column("apariencia", "logo_url", "TEXT")

    db.session.commit()


def init_db():
    with app.app_context():
        db.create_all()
        ensure_schema_updates()

        usuarios_iniciales = [
            ("Rosa Isela Vázquez Medina", "ROUSSVAM", "lic.rosavazquezm@gmail.com", "RIVaM061", "Administrador", "Administración"),
            ("Lic. Carmen Acosta C.", "lic.carmenacostac", "mrsmunoz68@gmail.com", "LIC.CARMENAC68", "Usuario", "Jurídico"),
            ("Computadora Oficina", "oficina.alesi", "", "Oficina2026*", "Usuario Oficina", "Oficina"),
        ]

        for nombre, usuario, correo, password, rol, area in usuarios_iniciales:
            existente = Usuario.query.filter_by(usuario=usuario).first()
            if not existente:
                db.session.add(Usuario(nombre=nombre, usuario=usuario, correo=correo, password_hash=generate_password_hash(password), rol=rol, area=area))
            else:
                existente.area = existente.area or area

        for materia, nombre in TIPOS_INICIALES:
            if not TipoAsunto.query.filter_by(materia=materia, nombre=nombre).first():
                db.session.add(TipoAsunto(materia=materia, nombre=nombre))

        if not Apariencia.query.first():
            db.session.add(Apariencia())

        db.session.commit()


def actual():
    if not session.get("usuario_id"):
        return None
    return db.session.get(Usuario, session["usuario_id"])


def req():
    return actual() is not None


def admin():
    u = actual()
    return bool(u and u.rol == "Administrador")


def subir(archivo):
    if not archivo or not archivo.filename:
        return ""
    if cloudinary and os.getenv("CLOUDINARY_CLOUD_NAME") and os.getenv("CLOUDINARY_API_KEY") and os.getenv("CLOUDINARY_API_SECRET"):
        cloudinary.config(
            cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
            api_key=os.getenv("CLOUDINARY_API_KEY"),
            api_secret=os.getenv("CLOUDINARY_API_SECRET"),
            secure=True,
        )
        resultado = cloudinary.uploader.upload(archivo, resource_type="auto", folder="alesi_grupo_juridico")
        return resultado.get("secure_url", "")
    nombre_seguro = secure_filename(archivo.filename)
    ruta = "uploads/" + datetime.now().strftime("%Y%m%d%H%M%S_") + nombre_seguro
    archivo.save(ruta)
    return "/" + ruta


def registrar_bitacora(expediente_id, accion, detalle=""):
    if actual():
        db.session.add(Bitacora(expediente_id=expediente_id, usuario_id=actual().id, accion=accion, detalle=detalle))


def permiso_expediente(expediente_id):
    # Todos los usuarios activos del despacho pueden ver y trabajar los expedientes.
    # Se mantiene esta función para no romper rutas anteriores.
    return "Administración" if actual() else ""


def puede_ver(expediente_id):
    return req()


def puede_editar(expediente_id):
    return req()


def puede_administrar(expediente_id):
    return req()


def usuarios_acceso(expediente_id):
    # Los avisos de expedientes se muestran a todo el despacho.
    return Usuario.query.filter(Usuario.activo == True).order_by(Usuario.nombre).all()


def correo(destinatario, asunto, cuerpo):
    if not all([os.getenv("SMTP_HOST"), os.getenv("SMTP_USER"), os.getenv("SMTP_PASS"), destinatario]):
        print("[SIMULADO]", destinatario, asunto)
        return False
    mensaje = EmailMessage()
    mensaje["Subject"] = asunto
    mensaje["From"] = os.getenv("FROM_EMAIL", os.getenv("SMTP_USER"))
    mensaje["To"] = destinatario
    mensaje.set_content(cuerpo)
    with smtplib.SMTP(os.getenv("SMTP_HOST"), int(os.getenv("SMTP_PORT", "587"))) as servidor:
        servidor.starttls()
        servidor.login(os.getenv("SMTP_USER"), os.getenv("SMTP_PASS"))
        servidor.send_message(mensaje)
    return True


def notificar_interna(usuario_id, titulo, mensaje, enlace="", creador_id=None, grupo=None):
    db.session.add(NotificacionInterna(
        usuario_id=usuario_id,
        titulo=titulo,
        mensaje=mensaje,
        enlace=enlace,
        creador_id=creador_id,
        grupo=grupo
    ))


def crear_aviso_para_accesos(expediente_id, titulo, mensaje):
    expediente = db.session.get(Expediente, expediente_id)
    if not expediente:
        return
    db.session.add(AvisoExpediente(expediente_id=expediente_id, usuario_id=actual().id if actual() else None, titulo=titulo, mensaje=mensaje, fecha_aviso=local_date_iso()))
    for usuario in usuarios_acceso(expediente_id):
        notificar_interna(usuario.id, titulo, mensaje, f"/expediente/{expediente_id}")


def revisar_y_enviar_notificaciones():
    enviados = 0
    ahora = datetime.now()
    ventana_minutos = 12
    items = []
    for audiencia in Audiencia.query.filter_by(notificar=True).all():
        items.append(("audiencia", audiencia.id, audiencia.expediente_id, audiencia.fecha, audiencia.hora, f"Audiencia: {audiencia.titulo}", audiencia.expediente.numero if audiencia.expediente else ""))
    for movimiento in Movimiento.query.filter_by(notificar=True).all():
        if movimiento.fecha_limite:
            items.append(("vencimiento", movimiento.id, movimiento.expediente_id, movimiento.fecha_limite, movimiento.hora_limite or "09:00", f"Vencimiento: {movimiento.titulo}", movimiento.expediente.numero if movimiento.expediente else ""))
    for tipo, referencia_id, expediente_id, fecha, hora, titulo, numero in items:
        try:
            fecha_hora = datetime.strptime(fecha + " " + hora, "%Y-%m-%d %H:%M")
        except Exception:
            continue
        for etiqueta, horas in [("24h", 24), ("2h", 2)]:
            momento_aviso = fecha_hora - timedelta(hours=horas)
            diferencia = abs((ahora - momento_aviso).total_seconds()) / 60
            if diferencia <= ventana_minutos:
                for usuario in usuarios_acceso(expediente_id):
                    ya_enviada = NotificacionEnviada.query.filter_by(tipo=tipo, referencia_id=referencia_id, anticipacion=etiqueta, usuario_id=usuario.id).first()
                    if ya_enviada:
                        continue
                    notificar_interna(usuario.id, f"{titulo} en {etiqueta}", f"Expediente {numero}. Fecha: {fecha} {hora}", f"/expediente/{expediente_id}")
                    if usuario.correo:
                        correo(usuario.correo, f"ALESI - {titulo} en {etiqueta}", f"Expediente: {numero}\n{titulo}\nFecha: {fecha} {hora}\n\nALESI GRUPO JURÍDICO")
                    db.session.add(NotificacionEnviada(tipo=tipo, referencia_id=referencia_id, anticipacion=etiqueta, usuario_id=usuario.id))
                    enviados += 1
    db.session.commit()
    return enviados


BASE = """
<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>ALESI GRUPO JURÍDICO</title>
<style>
:root{--n:{{apariencia.color_principal}};--t:{{apariencia.color_secundario}};--d:{{apariencia.color_dorado}}}
body{margin:0;font-family:{{apariencia.fuente}},Arial;background:{{apariencia.fondo}};color:#374151}
.top{background:var(--n);color:white;padding:8px 28px;display:flex;justify-content:space-between;gap:20px;align-items:center}
.clock{font-weight:bold;color:#f8fafc}
.hero{background:linear-gradient(120deg,#111827,#0f766e);color:white;padding:28px 36px;border-bottom:6px solid var(--d);display:flex;align-items:center}
.logo{width:100px;height:100px;border-radius:50%;background:white;object-fit:cover;border:4px solid var(--d);margin-right:20px}
nav{background:#0f172a;padding:12px 30px;line-height:2.2}
nav a{color:white;text-decoration:none;margin:4px 5px 4px 0;font-weight:bold;background:#1f2937;border:1px solid #334155;border-radius:999px;padding:9px 13px;display:inline-block} nav a:hover{background:var(--d);color:white}
main{padding:25px 30px}
.card{background:white;border-radius:14px;padding:22px;margin-bottom:20px;box-shadow:0 3px 12px #0002;border-top:4px solid var(--d)}
input,select,textarea{width:100%;padding:10px;margin:6px 0 12px;border:1px solid #cbd5e1;border-radius:8px;box-sizing:border-box}
button,.btn{background:var(--d);color:white;border:0;border-radius:8px;padding:10px 16px;text-decoration:none;font-weight:bold;display:inline-block;cursor:pointer}
.btn2{background:var(--t)}.btnDark{background:var(--n)}
table{width:100%;border-collapse:collapse} th,td{padding:10px;border-bottom:1px solid #e5e7eb;text-align:left;vertical-align:top} th{background:#e5e7eb}
.grid{display:grid;grid-template-columns:repeat(4,1fr);gap:15px}.grid2{display:grid;grid-template-columns:repeat(2,1fr);gap:15px}
.stat{text-align:center;background:white;border-radius:14px;padding:20px;box-shadow:0 3px 12px #0002;border-bottom:4px solid var(--t)}
.flash{background:#dcfce7;padding:10px;border-left:5px solid #166534;margin-bottom:12px}.badge{background:var(--t);color:white;padding:4px 8px;border-radius:999px;font-size:12px;font-weight:bold}
.small{font-size:13px;color:#6b7280}.chat{background:#f8fafc;border-left:4px solid var(--t);padding:10px;margin-bottom:8px;border-radius:8px}.bit{background:#fff7ed;border-left:4px solid var(--d);padding:9px;margin-bottom:7px;border-radius:8px}
@media print{nav,.top,button,.btn,.no-print,.no-imprimir{display:none!important}.card{box-shadow:none;border:0}.hero{display:none!important}main{padding:0}body{background:white}.print-list{display:block!important}.solo-imprimir{display:block!important}}
@media(max-width:900px){.grid,.grid2{grid-template-columns:1fr}}
</style>
<script>
function relojAlesi(){
    const el=document.getElementById("clock");
    if(!el) return;
    const d=new Date();
    el.textContent=d.toLocaleDateString("es-MX",{weekday:"long",year:"numeric",month:"long",day:"numeric"})+" | "+d.toLocaleTimeString("es-MX");
}
setInterval(relojAlesi,1000);
window.onload=relojAlesi;
</script>
</head>
<body>
<div class="top">
    <div>Sistema Jurídico + Planeador + Cobranza</div>
    <div class="clock" id="clock"></div>
    <div style="display:flex;align-items:center;gap:10px">{% if usuario %}{% if usuario.foto_url %}<img src="{{usuario.foto_url}}" style="width:34px;height:34px;border-radius:50%;object-fit:cover;border:2px solid #B38B2E">{% endif %}<span>{{usuario.nombre}} | {{usuario.rol}}</span><a class="btn btn2" href="/perfil">Mi Perfil</a><a class="btn btnDark" href="/logout">Cerrar sesión</a>{% endif %}</div>
</div>
<section class="hero">
    <img class="logo" src="/logo">
    <div><h1>ALESI GRUPO JURÍDICO</h1><p>Sistema Integral de Gestión Jurídica y Control de Expedientes</p></div>
</section>
{% if usuario %}
<nav>
    <a href="/">Inicio</a>
    <a href="/mis-expedientes">Mis Expedientes</a>
    <a href="/planeador">Planeador</a>
    <a href="/clientes">Clientes</a>
    <a href="/contratos">Contratos</a>
    <a href="/citas">Citas</a>
    <a href="/audiencias">Audiencias</a>
    <a href="/vencimientos">Vencimientos</a>
    <a href="/cobranza">Cobranza</a>
    {% if usuario.rol == 'Administrador' %}<a href="/respaldos">Respaldos</a>{% endif %}
    <a href="/notificaciones-internas">Avisos</a>
    <a href="/jurisprudencia">Jurisprudencia</a>
    <a href="/boletin-judicial">Boletín Judicial</a>
    <a href="/teja">TEJA</a>
    {% if usuario.rol == 'Administrador' %}
        <a href="/usuarios">Usuarios</a>
        <a href="/catalogo-asuntos">Catálogo</a>
        <a href="/notificaciones">Notificaciones</a>
        <a href="/apariencia">Apariencia</a>
    {% endif %}
</nav>
{% endif %}
<main>
{% with messages=get_flashed_messages() %}{% for m in messages %}<div class="flash">{{m}}</div>{% endfor %}{% endwith %}
{{contenido|safe}}
</main>
</body>
</html>
"""


def apariencia_actual():
    try:
        a = Apariencia.query.first()
        return a or Apariencia()
    except Exception:
        return Apariencia()


def render(contenido, **kw):
    return render_template_string(BASE, contenido=render_template_string(contenido, **kw), usuario=actual(), apariencia=apariencia_actual())


@app.route("/logo")
def logo():
    try:
        a = Apariencia.query.first()
        if a and a.logo_url:
            return redirect(a.logo_url)
    except Exception:
        pass
    if os.path.exists("logo_alesi.png"):
        return send_from_directory(".", "logo_alesi.png")
    return ("", 404)


@app.route("/uploads/<path:archivo>")
def uploads(archivo):
    return send_from_directory("uploads", archivo)


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        usuario = Usuario.query.filter_by(usuario=request.form["usuario"], activo=True).first()
        if usuario and check_password_hash(usuario.password_hash, request.form["password"]):
            session["usuario_id"] = usuario.id
            return redirect("/")
        flash("Usuario o contraseña incorrectos.")
    return render("""<div class="card" style="max-width:420px;margin:auto"><h2>Iniciar sesión</h2><form method="post"><label>Usuario</label><input name="usuario" required><label>Contraseña</label><input type="password" name="password" required><button>Entrar</button></form></div>""")


@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")


@app.route("/perfil", methods=["GET", "POST"])
def perfil():
    if not req():
        return redirect("/login")
    u = actual()
    if request.method == "POST":
        u.nombre = request.form.get("nombre", u.nombre)
        u.correo = request.form.get("correo", u.correo)
        nueva = request.form.get("password", "").strip()
        if nueva:
            u.password_hash = generate_password_hash(nueva)
        foto = subir(request.files.get("foto"))
        if foto:
            u.foto_url = foto
        db.session.commit()
        flash("Perfil actualizado.")
        return redirect("/perfil")
    return render("""<div class="card" style="max-width:760px;margin:auto"><h2>Mi Perfil</h2>{% if u.foto_url %}<p><img src="{{u.foto_url}}" style="width:120px;height:120px;border-radius:50%;object-fit:cover;border:3px solid #B38B2E"></p>{% endif %}<form method="post" enctype="multipart/form-data"><label>Nombre</label><input name="nombre" value="{{u.nombre or ''}}" required><label>Usuario</label><input value="{{u.usuario}}" disabled><label>Correo</label><input type="email" name="correo" value="{{u.correo or ''}}"><label>Nueva contraseña</label><input type="password" name="password" placeholder="Déjalo vacío si no deseas cambiarla"><label>Cambiar fotografía</label><input type="file" name="foto"><button>Guardar cambios</button></form></div>""", u=u)


@app.route("/")
def inicio():
    if not req():
        return redirect("/login")
    hoy = local_now().date()
    ayer = hoy - timedelta(days=1)
    manana = hoy + timedelta(days=1)
    fechas = [ayer.strftime("%Y-%m-%d"), hoy.strftime("%Y-%m-%d"), manana.strftime("%Y-%m-%d")]
    labels = {ayer.strftime("%Y-%m-%d"): "Ayer", hoy.strftime("%Y-%m-%d"): "Hoy", manana.strftime("%Y-%m-%d"): "Mañana"}

    pendientes = []
    for a in Audiencia.query.filter(Audiencia.fecha.in_(fechas)).order_by(Audiencia.fecha, Audiencia.hora).all():
        pendientes.append({"fecha": a.fecha, "label": labels.get(a.fecha, ""), "tipo": "Audiencia", "texto": f"{a.hora} - {a.expediente.numero if a.expediente else ''} - {a.titulo}"})
    for v in Movimiento.query.filter(Movimiento.fecha_limite.in_(fechas)).order_by(Movimiento.fecha_limite, Movimiento.hora_limite).all():
        pendientes.append({"fecha": v.fecha_limite, "label": labels.get(v.fecha_limite, ""), "tipo": "Vencimiento", "texto": f"{v.hora_limite or ''} - {v.expediente.numero if v.expediente else ''} - {v.proxima_accion or v.titulo}"})
    for p in Cobranza.query.filter(Cobranza.proximo_pago.in_(fechas), Cobranza.saldo > 0).order_by(Cobranza.proximo_pago).all():
        pendientes.append({"fecha": p.proximo_pago, "label": labels.get(p.proximo_pago, ""), "tipo": "Pago", "texto": f"{p.cliente.nombre if p.cliente else ''} - {p.moneda} {p.saldo}"})

    inicio_semana = hoy - timedelta(days=hoy.weekday())
    fin_semana = inicio_semana + timedelta(days=6)
    cobranza_semana = Cobranza.query.filter(Cobranza.proximo_pago >= inicio_semana.strftime("%Y-%m-%d"), Cobranza.proximo_pago <= fin_semana.strftime("%Y-%m-%d"), Cobranza.saldo > 0).count()
    avisos = NotificacionInterna.query.filter_by(usuario_id=actual().id, leida=False).count()

    return render("""<div class="grid"><div class="stat"><h2>{{clientes}}</h2><p>Clientes</p></div><div class="stat"><h2>{{expedientes}}</h2><p>Expedientes</p></div><div class="stat"><h2>{{audiencias}}</h2><p>Audiencias</p></div><div class="stat"><h2>{{vencimientos}}</h2><p>Vencimientos</p></div></div><br><div class="grid"><div class="stat"><h2>{{avisos}}</h2><p>Avisos internos no leídos</p></div><div class="stat"><h2>{{cobranza_semana}}</h2><p>Cobranza pendiente esta semana</p></div></div><br><div class="card"><h2>Pendientes de ayer, hoy y mañana</h2><table><tr><th>Día</th><th>Fecha</th><th>Tipo</th><th>Detalle</th></tr>{% for p in pendientes %}<tr><td>{{p.label}}</td><td>{{p.fecha}}</td><td>{{p.tipo}}</td><td>{{p.texto}}</td></tr>{% endfor %}</table></div>""", clientes=Cliente.query.count(), expedientes=Expediente.query.count(), audiencias=Audiencia.query.count(), vencimientos=Movimiento.query.filter(Movimiento.fecha_limite != "").count(), avisos=avisos, cobranza_semana=cobranza_semana, pendientes=pendientes)


@app.route("/usuarios", methods=["GET", "POST"])
def usuarios():
    if not req() or not admin():
        return redirect("/")
    if request.method == "POST":
        if Usuario.query.count() >= 15:
            flash("Límite de 15 usuarios.")
            return redirect("/usuarios")
        db.session.add(Usuario(nombre=request.form["nombre"], usuario=request.form["usuario"], correo=request.form["correo"], rol=request.form["rol"], area=request.form["area"], foto_url=subir(request.files.get("foto")), password_hash=generate_password_hash(request.form["password"])))
        db.session.commit()
        flash("Usuario creado.")
        return redirect("/usuarios")
    return render("""<div class="grid2"><div class="card"><h2>Crear usuario</h2><form method="post" enctype="multipart/form-data"><label>Nombre</label><input name="nombre" required><label>Usuario</label><input name="usuario" required><label>Correo</label><input type="email" name="correo"><label>Contraseña</label><input type="password" name="password" required><label>Rol</label><select name="rol"><option>Usuario</option><option>Administrador</option><option>Usuario Oficina</option></select><label>Área</label><select name="area"><option>Jurídico</option><option>Administración</option><option>Oficina</option><option>Penal</option><option>Laboral</option><option>Administrativo</option><option>Civil</option><option>Mercantil</option><option>Familiar</option></select><label>Foto</label><input type="file" name="foto"><button>Guardar</button></form></div><div class="card"><h2>Usuarios</h2><table><tr><th>Nombre</th><th>Usuario</th><th>Correo</th><th>Rol</th><th>Área</th></tr>{% for u in datos %}<tr><td>{{u.nombre}}</td><td>{{u.usuario}}</td><td>{{u.correo}}</td><td>{{u.rol}}</td><td>{{u.area}}</td></tr>{% endfor %}</table></div></div>""", datos=Usuario.query.order_by(Usuario.id).all())


@app.route("/catalogo-asuntos", methods=["GET", "POST"])
def catalogo_asuntos():
    if not req() or not admin():
        return redirect("/")
    if request.method == "POST":
        db.session.add(TipoAsunto(materia=request.form["materia"], nombre=request.form["nombre"]))
        db.session.commit()
        flash("Tipo de asunto agregado.")
        return redirect("/catalogo-asuntos")
    return render("""<div class="grid2"><div class="card"><h2>Agregar tipo de asunto</h2><form method="post"><label>Materia</label><select name="materia">{% for m in materias %}<option>{{m}}</option>{% endfor %}</select><label>Tipo de asunto</label><input name="nombre" required><button>Guardar</button></form></div><div class="card"><h2>Catálogo de asuntos</h2><table><tr><th>Materia</th><th>Tipo de asunto</th></tr>{% for t in tipos %}<tr><td>{{t.materia}}</td><td>{{t.nombre}}</td></tr>{% endfor %}</table></div></div>""", materias=MATERIAS, tipos=TipoAsunto.query.order_by(TipoAsunto.materia, TipoAsunto.nombre).all())


@app.route("/clientes", methods=["GET", "POST"])
def clientes():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        db.session.add(Cliente(
            nombre=request.form["nombre"], telefono=request.form.get("telefono", ""), correo=request.form.get("correo", ""),
            rfc=request.form.get("rfc", ""), curp=request.form.get("curp", ""), direccion=request.form.get("direccion", ""),
            observaciones=request.form.get("observaciones", "")
        ))
        db.session.commit()
        flash("Cliente guardado.")
        return redirect("/clientes")
    return render("""<div class="card no-print no-imprimir"><h2>Registrar cliente</h2><form method="post"><label>Nombre</label><input name="nombre" required><label>Teléfono</label><input name="telefono"><label>Correo</label><input name="correo"><label>RFC</label><input name="rfc"><label>CURP</label><input name="curp"><label>Dirección</label><textarea name="direccion"></textarea><label>Observaciones</label><textarea name="observaciones"></textarea><button>Guardar</button></form></div><div class="card print-list"><h2>Lista de clientes <button class="no-print" onclick="window.print()">Imprimir lista</button></h2><table><tr><th>Cliente</th><th>Teléfono</th><th>Correo</th><th>RFC</th><th class="no-print">Acciones</th></tr>{% for c in datos %}<tr><td><a href="/cliente/{{c.id}}">{{c.nombre}}</a></td><td>{{c.telefono}}</td><td>{{c.correo}}</td><td>{{c.rfc}}</td><td class="no-print"><a class="btn btn2" href="/editar-cliente/{{c.id}}">Editar</a> <a class="btn btnDark" href="/eliminar-cliente/{{c.id}}" onclick="return confirm('¿Eliminar cliente? Se ocultará de la lista, pero sus expedientes se conservan.')">Eliminar</a></td></tr>{% endfor %}</table></div>""", datos=Cliente.query.filter(Cliente.activo == True).order_by(Cliente.nombre).all())


@app.route("/cliente/<int:id>")
def cliente(id):
    if not req():
        return redirect("/login")
    c = Cliente.query.get_or_404(id)
    pagos = Cobranza.query.filter_by(cliente_id=id).order_by(Cobranza.id.desc()).all()
    exps = Expediente.query.filter_by(cliente_id=id).filter(Expediente.estado != "Eliminado").order_by(Expediente.id.desc()).all()
    resumen = []
    for e in exps:
        abonado = db.session.query(db.func.sum(Cobranza.abono)).filter_by(expediente_id=e.id).scalar() or 0
        total = e.cobro_total or 0
        resumen.append({"e": e, "abonado": abonado, "saldo": max(total - abonado, 0)})
    return render("""<div class="card no-print"><h2>Cliente</h2><p><a class="btn btn2" href="/editar-cliente/{{c.id}}">Editar datos</a> <a class="btn" href="/contrato-cliente/{{c.id}}">Crear contrato PDF</a> <button onclick="window.print()">Imprimir datos del cliente</button></p></div><div class="card print-list"><h2>Datos del cliente</h2><p><b>Nombre:</b> {{c.nombre}}</p><p><b>Teléfono:</b> {{c.telefono}}</p><p><b>Correo:</b> {{c.correo}}</p><p><b>RFC:</b> {{c.rfc}}</p><p><b>CURP:</b> {{c.curp}}</p><p><b>Dirección:</b> {{c.direccion}}</p><p>{{c.observaciones}}</p></div><div class="card print-list"><h2>Expedientes iniciados o concluidos</h2><table><tr><th>Número</th><th>Materia</th><th>Tipo</th><th>Estado</th><th>Cobro</th><th>Abonado</th><th>Saldo</th><th class="no-print">Abrir</th></tr>{% for r in resumen %}<tr><td>{{r.e.numero}}</td><td>{{r.e.materia}}</td><td>{{r.e.tipo_asunto}}</td><td>{{r.e.estado}}</td><td>{{r.e.moneda_cobro}} {{'%.2f'|format(r.e.cobro_total or 0)}}</td><td>{{'%.2f'|format(r.abonado or 0)}}</td><td>{{'%.2f'|format(r.saldo or 0)}}</td><td class="no-print"><a class="btn btn2" href="/expediente/{{r.e.id}}">Abrir</a></td></tr>{% endfor %}</table></div><div class="card print-list"><h2>Pagos realizados desglosados por expediente</h2><table><tr><th>Fecha</th><th>Folio</th><th>Expediente</th><th>Concepto</th><th>Abono</th><th>Saldo</th><th>Recibo</th></tr>{% for p in pagos %}<tr><td>{{p.fecha_pago}}</td><td>{{p.folio}}</td><td>{{p.expediente.numero if p.expediente else ''}}</td><td>{{p.concepto}}</td><td>{{p.moneda}} {{'%.2f'|format(p.abono or 0)}}</td><td>{{'%.2f'|format(p.saldo or 0)}}</td><td class="no-print"><a href="/recibo/{{p.id}}">Ver</a> | <a href="/recibo-pdf/{{p.id}}">PDF</a></td></tr>{% endfor %}</table></div>""", c=c, pagos=pagos, exps=exps, resumen=resumen)


@app.route("/editar-cliente/<int:id>", methods=["GET", "POST"])
def editar_cliente(id):
    if not req():
        return redirect("/login")
    c = Cliente.query.get_or_404(id)
    if request.method == "POST":
        c.nombre = request.form.get("nombre", c.nombre)
        c.telefono = request.form.get("telefono", "")
        c.correo = request.form.get("correo", "")
        c.rfc = request.form.get("rfc", "")
        c.curp = request.form.get("curp", "")
        c.direccion = request.form.get("direccion", "")
        c.observaciones = request.form.get("observaciones", "")
        db.session.commit()
        flash("Cliente actualizado.")
        return redirect("/cliente/" + str(id))
    return render("""<div class="card"><h2>Editar cliente</h2><form method="post"><label>Nombre</label><input name="nombre" value="{{c.nombre or ''}}" required><label>Teléfono</label><input name="telefono" value="{{c.telefono or ''}}"><label>Correo</label><input name="correo" value="{{c.correo or ''}}"><label>RFC</label><input name="rfc" value="{{c.rfc or ''}}"><label>CURP</label><input name="curp" value="{{c.curp or ''}}"><label>Dirección</label><textarea name="direccion">{{c.direccion or ''}}</textarea><label>Observaciones</label><textarea name="observaciones">{{c.observaciones or ''}}</textarea><button>Guardar cambios</button></form></div>""", c=c)


@app.route("/eliminar-cliente/<int:id>")
def eliminar_cliente(id):
    if not req():
        return redirect("/login")
    c = Cliente.query.get_or_404(id)
    c.activo = False
    db.session.commit()
    flash("Cliente eliminado de la lista. Sus expedientes y pagos se conservaron.")
    return redirect("/clientes")


@app.route("/expedientes", methods=["GET", "POST"])
def expedientes():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        nuevo = Expediente(
            propietario_id=actual().id,
            cliente_id=request.form.get("cliente_id") or None,
            numero=request.form["numero"],
            materia=request.form["materia"],
            tipo_asunto=request.form["tipo_asunto"],
            cobro_total=float(request.form.get("cobro_total") or 0),
            moneda_cobro=request.form.get("moneda_cobro", "MXN"),
            abono_pactado=float(request.form.get("abono_pactado") or 0),
            periodicidad_abono=request.form.get("periodicidad_abono", ""),
            autoridad=request.form.get("autoridad", ""),
            actor=request.form.get("actor", ""),
            demandado=request.form.get("demandado", ""),
            estado=request.form.get("estado", "En trámite"),
            prioridad=request.form.get("prioridad", "Media"),
            responsable=request.form.get("responsable", ""),
            fecha_inicio=request.form.get("fecha_inicio", ""),
            observaciones=request.form.get("observaciones", ""),
        )
        db.session.add(nuevo)
        db.session.flush()
        registrar_bitacora(nuevo.id, "Creó expediente", f"Expediente {nuevo.numero}")
        db.session.commit()
        flash("Expediente creado.")
        return redirect("/mis-expedientes")
    return redirect("/mis-expedientes")


@app.route("/mis-expedientes", methods=["GET", "POST"])
def mis_expedientes():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        return expedientes()
    datos = Expediente.query.filter(Expediente.estado != "Eliminado").order_by(Expediente.id.desc()).all()
    return render("""<div class="card no-print"><h2>Nuevo expediente</h2><form method="post" action="/expedientes"><label>Cliente</label><select name="cliente_id"><option value="">Sin cliente</option>{% for cliente in clientes %}<option value="{{cliente.id}}">{{cliente.nombre}}</option>{% endfor %}</select><label>Número</label><input name="numero" required><label>Materia</label><select name="materia">{% for m in materias %}<option>{{m}}</option>{% endfor %}</select><label>Tipo de asunto</label><select name="tipo_asunto">{% for t in tipos %}<option>{{t.nombre}}</option>{% endfor %}</select><label>Cantidad total que pagará</label><input type="number" step="0.01" name="cobro_total"><label>Moneda</label><select name="moneda_cobro"><option>MXN</option><option>USD</option></select><label>Cuánto debe abonar</label><input type="number" step="0.01" name="abono_pactado"><label>Periodicidad del abono</label><select name="periodicidad_abono"><option>Semanal</option><option>Quincenal</option><option>Mensual</option><option>Otro</option></select><label>Autoridad</label><input name="autoridad"><label>Actor</label><input name="actor"><label>Demandado</label><input name="demandado"><label>Estado</label><select name="estado"><option>Nuevo</option><option selected>En trámite</option><option>Pendiente de acuerdo</option><option>Pendiente de audiencia</option><option>Concluido</option><option>Archivado</option><option>Urgente</option></select><label>Prioridad</label><select name="prioridad"><option>Baja</option><option selected>Media</option><option>Alta</option></select><label>Responsable</label><input name="responsable"><label>Fecha inicio</label><input type="date" name="fecha_inicio"><label>Observaciones</label><textarea name="observaciones"></textarea><button>Guardar expediente</button></form></div><div class="card print-list"><h2>Expedientes del despacho <button class="no-print" onclick="window.print()">Imprimir lista</button></h2><p class="small">Todos los usuarios verán los expedientes registrados por el despacho.</p><table><tr><th>Número</th><th>Cliente</th><th>Materia</th><th>Tipo</th><th>Estado</th><th>Cobro</th><th>Abono</th><th>Periodicidad</th><th class="no-print"></th></tr>{% for e in datos %}<tr><td>{{e.numero}}</td><td>{{e.cliente.nombre if e.cliente else ""}}</td><td>{{e.materia}}</td><td>{{e.tipo_asunto}}</td><td>{{e.estado}}</td><td>{{e.moneda_cobro}} {{'%.2f'|format(e.cobro_total or 0)}}</td><td>{{'%.2f'|format(e.abono_pactado or 0)}}</td><td>{{e.periodicidad_abono}}</td><td class="no-print"><a class="btn btn2" href="/expediente/{{e.id}}">Abrir</a></td></tr>{% endfor %}</table></div>""", datos=datos, clientes=Cliente.query.filter(Cliente.activo == True).order_by(Cliente.nombre).all(), materias=MATERIAS, tipos=TipoAsunto.query.order_by(TipoAsunto.materia, TipoAsunto.nombre).all())


@app.route("/compartidos-conmigo")
def compartidos_conmigo():
    return redirect("/mis-expedientes")


@app.route("/expediente/<int:id>")
def expediente(id):
    if not req():
        return redirect("/login")
    if not puede_ver(id):
        flash("No tienes acceso a este expediente.")
        return redirect("/")
    expediente = Expediente.query.get_or_404(id)
    return render("""<div class="card"><h2>Expediente {{expediente.numero}} <button onclick="window.print()">Imprimir</button></h2><p><b>Cliente:</b> {{expediente.cliente.nombre if expediente.cliente else ""}}</p><p><b>Materia:</b> {{expediente.materia}}</p><p><b>Tipo de asunto:</b> {{expediente.tipo_asunto}}</p><p><b>Autoridad:</b> {{expediente.autoridad}}</p><p><b>Actor:</b> {{expediente.actor}}</p><p><b>Demandado:</b> {{expediente.demandado}}</p><p><b>Estado:</b> {{expediente.estado}}</p>{% if puede_editar %}<form method="post" action="/actualizar-estatus/{{expediente.id}}"><label>Cambiar estatus</label><select name="estado"><option {% if expediente.estado == 'Nuevo' %}selected{% endif %}>Nuevo</option><option {% if expediente.estado == 'En trámite' %}selected{% endif %}>En trámite</option><option {% if expediente.estado == 'Pendiente de acuerdo' %}selected{% endif %}>Pendiente de acuerdo</option><option {% if expediente.estado == 'Pendiente de audiencia' %}selected{% endif %}>Pendiente de audiencia</option><option {% if expediente.estado == 'Concluido' %}selected{% endif %}>Concluido</option><option {% if expediente.estado == 'Archivado' %}selected{% endif %}>Archivado</option><option {% if expediente.estado == 'Urgente' %}selected{% endif %}>Urgente</option></select><button>Actualizar estatus</button></form>{% endif %}<p><b>Tu permiso:</b> <span class="badge">{{permiso}}</span></p><p>{{expediente.observaciones}}</p>{% if puede_editar %}<a class="btn" href="/movimiento/{{expediente.id}}">Agregar promoción/vencimiento</a> <a class="btn btn2" href="/audiencia/{{expediente.id}}">Programar audiencia</a> <a class="btn" href="/mensaje/{{expediente.id}}">Mensaje interno</a>{% endif %} {% if puede_editar %}<a class="btn" href="/aviso-expediente/{{expediente.id}}">Agregar aviso</a>{% endif %}</div><div class="grid2"><div class="card"><h2>Promociones / vencimientos</h2><table><tr><th>Título</th><th>Usuario</th><th>Límite</th><th>Archivo</th></tr>{% for movimiento in movimientos %}<tr><td>{{movimiento.titulo}}</td><td>{{movimiento.usuario.nombre}}</td><td>{{movimiento.fecha_limite}} {{movimiento.hora_limite}}</td><td>{% if movimiento.archivo_url %}<a href="{{movimiento.archivo_url}}" target="_blank">Ver</a>{% endif %}</td></tr>{% endfor %}</table></div><div class="card"><h2>Audiencias</h2><table><tr><th>Fecha</th><th>Hora</th><th>Audiencia</th><th>Usuario</th></tr>{% for audiencia in audiencias %}<tr><td>{{audiencia.fecha}}</td><td>{{audiencia.hora}}</td><td>{{audiencia.titulo}}</td><td>{{audiencia.usuario.nombre}}</td></tr>{% endfor %}</table></div></div><div class="grid2"><div class="card"><h2>Chat interno</h2>{% for m in mensajes %}<div class="chat"><b>{{m.usuario.nombre}}</b><div class="small">{{m.creado_en.strftime("%d/%m/%Y %H:%M")}}</div><p>{{m.mensaje}}</p></div>{% endfor %}</div><div class="card"><h2>Avisos del expediente</h2>{% for av in avisos_exp %}<div class="chat"><b>{{av.titulo}}</b> - {{av.creado_en.strftime('%d/%m/%Y %H:%M')}}<br>{{av.mensaje}}</div>{% endfor %}</div></div><div class="card"><h2>Bitácora</h2>{% for b in bitacora %}<div class="bit"><b>{{b.accion}}</b><div class="small">{{b.usuario.nombre if b.usuario else ""}} | {{b.creado_en.strftime("%d/%m/%Y %H:%M")}}</div><p>{{b.detalle}}</p></div>{% endfor %}</div>""", expediente=expediente, permiso=permiso_expediente(id), puede_editar=puede_editar(id), puede_administrar=puede_administrar(id), movimientos=Movimiento.query.filter_by(expediente_id=id).order_by(Movimiento.id.desc()).all(), audiencias=Audiencia.query.filter_by(expediente_id=id).order_by(Audiencia.fecha, Audiencia.hora).all(), mensajes=MensajeExpediente.query.filter_by(expediente_id=id).order_by(MensajeExpediente.id.desc()).all(), bitacora=Bitacora.query.filter_by(expediente_id=id).order_by(Bitacora.id.desc()).all(), avisos_exp=AvisoExpediente.query.filter_by(expediente_id=id).order_by(AvisoExpediente.id.desc()).all())

@app.route("/eliminar-expediente/<int:id>")
def eliminar_expediente(id):
    if not req():
        return redirect("/login")
    expediente = Expediente.query.get_or_404(id)
    if not (admin() or expediente.propietario_id == actual().id):
        flash("No tienes permiso para eliminar este expediente.")
        return redirect("/expediente/" + str(id))
    expediente.estado = "Eliminado"
    registrar_bitacora(id, "Expediente ocultado", "Se ocultó de la lista, pero no se borró de la base de datos.")
    db.session.commit()
    flash("Expediente ocultado correctamente. No se borró definitivamente.")
    return redirect("/mis-expedientes")


@app.route("/actualizar-estatus/<int:id>", methods=["POST"])
def actualizar_estatus(id):
    if not req():
        return redirect("/login")

    if not puede_editar(id):
        flash("No tienes permiso para cambiar el estatus.")
        return redirect("/expediente/" + str(id))

    expediente = Expediente.query.get_or_404(id)
    nuevo_estatus = request.form["estado"]

    expediente.estado = nuevo_estatus
    registrar_bitacora(id, "Cambió estatus", f"Nuevo estatus: {nuevo_estatus}")
    crear_aviso_para_accesos(id, "Estatus actualizado", f"El expediente {expediente.numero} cambió a: {nuevo_estatus}")

    db.session.commit()
    flash("Estatus actualizado.")
    return redirect("/expediente/" + str(id))

@app.route("/mensaje/<int:expediente_id>", methods=["GET", "POST"])
def mensaje(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_editar(expediente_id):
        flash("No tienes permiso para agregar mensajes.")
        return redirect("/expediente/" + str(expediente_id))
    if request.method == "POST":
        texto = request.form["mensaje"]
        db.session.add(MensajeExpediente(expediente_id=expediente_id, usuario_id=actual().id, mensaje=texto))
        registrar_bitacora(expediente_id, "Agregó mensaje interno", texto[:120])
        crear_aviso_para_accesos(expediente_id, "Nuevo mensaje interno", f"{actual().nombre}: {texto[:220]}")
        db.session.commit()
        flash("Mensaje agregado y avisado.")
        return redirect("/expediente/" + str(expediente_id))
    return render("""<div class="card"><h2>Mensaje interno</h2><form method="post"><label>Mensaje</label><textarea name="mensaje" required></textarea><button>Enviar</button></form></div>""")

@app.route("/eliminar-aviso/<int:id>")
def eliminar_aviso(id):
    if not req():
        return redirect("/login")

    aviso = AvisoExpediente.query.get_or_404(id)

    if not (admin() or aviso.usuario_id == actual().id):
        flash("No tienes permiso para eliminar este aviso.")
        return redirect("/expediente/" + str(aviso.expediente_id))

    expediente_id = aviso.expediente_id

    db.session.delete(aviso)
    db.session.commit()

    flash("Aviso eliminado.")
    return redirect("/expediente/" + str(expediente_id))
    
@app.route("/aviso-expediente/<int:expediente_id>", methods=["GET", "POST"])
def aviso_expediente(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_editar(expediente_id):
        flash("No tienes permiso para agregar avisos.")
        return redirect("/expediente/" + str(expediente_id))
    if request.method == "POST":
        titulo = request.form["titulo"]
        mensaje = request.form["mensaje"]
        crear_aviso_para_accesos(expediente_id, titulo, mensaje)
        registrar_bitacora(expediente_id, "Aviso agregado", titulo)
        db.session.commit()
        flash("Aviso agregado.")
        return redirect("/expediente/" + str(expediente_id))
    return render("""<div class="card"><h2>Agregar aviso al expediente</h2><form method="post"><label>Título</label><input name="titulo" required><label>Mensaje</label><textarea name="mensaje" required></textarea><button>Guardar aviso</button></form></div>""")


@app.route("/movimiento/<int:expediente_id>", methods=["GET", "POST"])
def movimiento(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_editar(expediente_id):
        flash("No tienes permiso para agregar movimientos.")
        return redirect("/expediente/" + str(expediente_id))
    if request.method == "POST":
        db.session.add(Movimiento(
            expediente_id=expediente_id, usuario_id=actual().id, titulo=request.form["titulo"], fecha=request.form["fecha"],
            estatus=request.form["estatus"], proxima_accion=request.form["proxima_accion"], fecha_limite=request.form["fecha_limite"],
            hora_limite=request.form["hora_limite"], observaciones=request.form["observaciones"],
            archivo_url=subir(request.files.get("archivo")), notificar=bool(request.form.get("notificar"))
        ))
        registrar_bitacora(expediente_id, "Agregó promoción/vencimiento", request.form["titulo"])
        crear_aviso_para_accesos(expediente_id, "Nuevo movimiento", f"{actual().nombre} agregó: {request.form['titulo']}")
        db.session.commit()
        flash("Guardado.")
        return redirect("/expediente/" + str(expediente_id))
    return render("""<div class="card"><h2>Agregar promoción / vencimiento</h2><form method="post" enctype="multipart/form-data"><label>Título</label><input name="titulo" required><label>Fecha</label><input type="date" name="fecha"><label>Estatus</label><select name="estatus"><option>Elaborado</option><option>Presentado</option><option>Acordado</option><option>Pendiente</option><option>Concluido</option></select><label>Próxima acción</label><input name="proxima_accion"><label>Fecha límite</label><input type="date" name="fecha_limite"><label>Hora límite</label><input type="time" name="hora_limite" value="09:00"><label><input type="checkbox" name="notificar" checked style="width:auto"> Notificar 24h y 2h antes</label><label>Observaciones</label><textarea name="observaciones"></textarea><label>Archivo</label><input type="file" name="archivo"><button>Guardar</button></form></div>""")


@app.route("/audiencia/<int:expediente_id>", methods=["GET", "POST"])
def audiencia(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_editar(expediente_id):
        flash("No tienes permiso para programar audiencias.")
        return redirect("/expediente/" + str(expediente_id))
    if request.method == "POST":
        db.session.add(Audiencia(
            expediente_id=expediente_id, usuario_id=actual().id, titulo=request.form["titulo"], fecha=request.form["fecha"],
            hora=request.form["hora"], autoridad=request.form["autoridad"], sala=request.form["sala"],
            modalidad=request.form["modalidad"], enlace=request.form["enlace"], observaciones=request.form["observaciones"],
            notificar=bool(request.form.get("notificar"))
        ))
        registrar_bitacora(expediente_id, "Programó audiencia", f'{request.form["titulo"]} - {request.form["fecha"]} {request.form["hora"]}')
        crear_aviso_para_accesos(expediente_id, "Audiencia programada", f"{request.form['titulo']} - {request.form['fecha']} {request.form['hora']}")
        db.session.commit()
        flash("Audiencia programada.")
        return redirect("/expediente/" + str(expediente_id))
    return render("""<div class="card"><h2>Programar audiencia</h2><form method="post"><label>Audiencia</label><input name="titulo" required><label>Fecha</label><input type="date" name="fecha" required><label>Hora</label><input type="time" name="hora" required><label>Cobro pactado para este asunto</label><input type="number" step="0.01" name="cobro_total"><label>Moneda</label><select name="moneda_cobro"><option>MXN</option><option>USD</option></select><label>Autoridad</label><input name="autoridad"><label>Sala</label><input name="sala"><label>Modalidad</label><select name="modalidad"><option>Presencial</option><option>Virtual</option><option>Mixta</option></select><label>Enlace</label><input name="enlace"><label><input type="checkbox" name="notificar" checked style="width:auto"> Notificar 24h y 2h antes</label><label>Observaciones</label><textarea name="observaciones"></textarea><button>Guardar</button></form></div>""")


@app.route("/compartir/<int:expediente_id>", methods=["GET", "POST"])
def compartir(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_administrar(expediente_id):
        flash("No tienes permiso para compartir este expediente.")
        return redirect("/expediente/" + str(expediente_id))
    expediente = Expediente.query.get_or_404(expediente_id)
    if request.method == "POST":
        usuario_id = int(request.form["usuario_id"])
        permiso = request.form["permiso"]
        existente = Compartido.query.filter_by(expediente_id=expediente_id, usuario_id=usuario_id).first()
        if existente:
            existente.permiso = permiso
            accion = "Actualizó permiso"
        else:
            db.session.add(Compartido(expediente_id=expediente_id, usuario_id=usuario_id, permiso=permiso))
            accion = "Compartió expediente"
        destino = db.session.get(Usuario, usuario_id)
        registrar_bitacora(expediente_id, accion, f"Usuario: {destino.nombre} | Permiso: {permiso}")
        notificar_interna(usuario_id, "Expediente compartido", f"{actual().nombre} compartió contigo el expediente {expediente.numero} con permiso {permiso}.", f"/expediente/{expediente_id}")
        if destino and destino.correo:
            correo(destino.correo, "ALESI - Expediente compartido", f"{actual().nombre} compartió contigo el expediente {expediente.numero}.\nPermiso: {permiso}\n\nALESI GRUPO JURÍDICO")
        db.session.commit()
        flash("Expediente compartido.")
        return redirect("/compartir/" + str(expediente_id))
    accesos = Compartido.query.filter_by(expediente_id=expediente_id).all()
    usuarios = Usuario.query.filter(Usuario.activo == True, Usuario.id != expediente.propietario_id).order_by(Usuario.nombre).all()
    return render("""<div class="grid2"><div class="card"><h2>Compartir expediente {{expediente.numero}}</h2><form method="post"><label>Usuario registrado</label><select name="usuario_id">{% for usuario in usuarios %}<option value="{{usuario.id}}">{{usuario.nombre}} - {{usuario.usuario}} - {{usuario.area}}</option>{% endfor %}</select><label>Permiso</label><select name="permiso"><option>Lectura</option><option>Edición</option><option>Administración</option></select><button>Compartir</button></form><br><a class="btn btn2" href="/compartir-area/{{expediente.id}}">Compartir por área</a></div><div class="card"><h2>Usuarios con acceso</h2><p><b>Propietario:</b> {{expediente.propietario.nombre}} <span class="badge">Administración</span></p><table><tr><th>Usuario</th><th>Permiso</th><th>Acción</th></tr>{% for a in accesos %}<tr><td>{{a.usuario.nombre}}</td><td><span class="badge">{{a.permiso}}</span></td><td><a class="btn btnDark" href="/quitar-acceso/{{a.id}}">Quitar</a></td></tr>{% endfor %}</table></div></div>""", expediente=expediente, usuarios=usuarios, accesos=accesos)


@app.route("/compartir-area/<int:expediente_id>", methods=["GET", "POST"])
def compartir_area(expediente_id):
    if not req():
        return redirect("/login")
    if not puede_administrar(expediente_id):
        flash("No tienes permiso para compartir por área.")
        return redirect("/expediente/" + str(expediente_id))
    expediente = Expediente.query.get_or_404(expediente_id)
    areas = sorted({u.area for u in Usuario.query.filter_by(activo=True).all() if u.area})
    if request.method == "POST":
        area = request.form["area"]
        permiso = request.form["permiso"]
        usuarios_area = Usuario.query.filter_by(area=area, activo=True).all()
        for usuario in usuarios_area:
            if usuario.id == expediente.propietario_id:
                continue
            existente = Compartido.query.filter_by(expediente_id=expediente_id, usuario_id=usuario.id).first()
            if existente:
                existente.permiso = permiso
            else:
                db.session.add(Compartido(expediente_id=expediente_id, usuario_id=usuario.id, permiso=permiso))
            notificar_interna(usuario.id, "Expediente compartido por área", f"Se compartió contigo el expediente {expediente.numero} por pertenecer al área {area}.", f"/expediente/{expediente_id}")
        registrar_bitacora(expediente_id, "Compartió expediente por área", f"Área: {area} | Permiso: {permiso}")
        db.session.commit()
        flash("Expediente compartido por área.")
        return redirect("/compartir/" + str(expediente_id))
    return render("""<div class="card"><h2>Compartir expediente por área</h2><form method="post"><label>Área</label><select name="area">{% for area in areas %}<option>{{area}}</option>{% endfor %}</select><label>Permiso</label><select name="permiso"><option>Lectura</option><option>Edición</option><option>Administración</option></select><button>Compartir por área</button></form></div>""", areas=areas)


@app.route("/quitar-acceso/<int:compartido_id>")
def quitar_acceso(compartido_id):
    if not req():
        return redirect("/login")
    compartido = Compartido.query.get_or_404(compartido_id)
    expediente_id = compartido.expediente_id
    if not puede_administrar(expediente_id):
        flash("No tienes permiso para quitar accesos.")
        return redirect("/expediente/" + str(expediente_id))
    nombre = compartido.usuario.nombre
    db.session.delete(compartido)
    registrar_bitacora(expediente_id, "Quitó acceso", f"Usuario: {nombre}")
    db.session.commit()
    flash("Acceso retirado.")
    return redirect("/compartir/" + str(expediente_id))


@app.route("/audiencias")
def audiencias():
    if not req():
        return redirect("/login")
    datos = Audiencia.query.order_by(Audiencia.fecha, Audiencia.hora).all()
    return render("""<div class="card"><h2>Audiencias <button onclick="window.print()">Imprimir</button></h2><table><tr><th>Fecha</th><th>Hora</th><th>Expediente</th><th>Audiencia</th><th>Autoridad</th></tr>{% for a in datos %}<tr><td>{{a.fecha}}</td><td>{{a.hora}}</td><td>{{a.expediente.numero}}</td><td>{{a.titulo}}</td><td>{{a.autoridad}}</td></tr>{% endfor %}</table></div>""", datos=datos)


@app.route("/vencimientos")
def vencimientos():
    if not req():
        return redirect("/login")
    datos = Movimiento.query.filter(Movimiento.fecha_limite != "").order_by(Movimiento.fecha_limite, Movimiento.hora_limite).all()
    return render("""<div class="card"><h2>Vencimientos <button onclick="window.print()">Imprimir</button></h2><table><tr><th>Fecha</th><th>Hora</th><th>Expediente</th><th>Movimiento</th><th>Acción</th></tr>{% for m in datos %}<tr><td>{{m.fecha_limite}}</td><td>{{m.hora_limite}}</td><td>{{m.expediente.numero}}</td><td>{{m.titulo}}</td><td>{{m.proxima_accion}}</td></tr>{% endfor %}</table></div>""", datos=datos)


@app.route("/cobranza", methods=["GET", "POST"])
def cobranza():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        cliente_id = request.form.get("cliente_id") or None
        expediente_id = request.form.get("expediente_id") or None
        expediente = db.session.get(Expediente, int(expediente_id)) if expediente_id else None
        cliente = db.session.get(Cliente, int(cliente_id)) if cliente_id else (expediente.cliente if expediente else None)
        total = float(request.form.get("monto_total") or (expediente.cobro_total if expediente else 0) or 0)
        abono = float(request.form.get("abono") or 0)
        pagos_previos = db.session.query(db.func.sum(Cobranza.abono)).filter_by(expediente_id=expediente.id).scalar() or 0 if expediente else 0
        saldo = max(total - pagos_previos - abono, 0)
        folio = "ALESI-" + local_stamp()
        p = Cobranza(
            folio=folio, cliente_id=cliente.id if cliente else cliente_id, expediente_id=expediente_id,
            usuario_id=actual().id, concepto=request.form.get("concepto", ""), descripcion=request.form.get("descripcion", ""),
            monto_total=total, abono=abono, saldo=saldo, moneda=request.form.get("moneda", "MXN"),
            forma_pago=request.form.get("forma_pago", "Efectivo"), fecha_pago=request.form.get("fecha_pago") or local_date_iso(),
            proximo_pago="" if saldo <= 0 else request.form.get("proximo_pago", ""), estatus="Liquidado" if saldo <= 0 else "Pendiente",
            comprobante_url=subir(request.files.get("comprobante"))
        )
        db.session.add(p)
        db.session.commit()
        flash("Cobranza guardada. Recibo generado.")
        return redirect("/recibo/" + str(p.id))
    expedientes = Expediente.query.filter(Expediente.estado != "Eliminado").order_by(Expediente.numero).all()
    return render("""<div class="card no-print no-imprimir"><h2>Registrar cobranza</h2><form method="post" enctype="multipart/form-data"><label>Cliente</label><select name="cliente_id"><option value="">Seleccione</option>{% for c in clientes %}<option value="{{c.id}}">{{c.nombre}}</option>{% endfor %}</select><label>Expediente</label><select name="expediente_id" onchange="llenarExpediente(this)"><option value="">Sin expediente</option>{% for e in expedientes %}<option value="{{e.id}}" data-total="{{e.cobro_total or 0}}" data-moneda="{{e.moneda_cobro or 'MXN'}}">{{e.numero}} - {{e.cliente.nombre if e.cliente else ''}} - {{e.tipo_asunto}} - {{e.moneda_cobro}} {{e.cobro_total}}</option>{% endfor %}</select><label>Concepto</label><select name="concepto"><option>Honorarios</option><option>Abono</option><option>Pago trámite institucional</option><option>Copias certificadas</option><option>Notario</option><option>Viáticos</option><option>Derechos</option><option>Gastos administrativos</option><option>Otro</option></select><label>Descripción</label><textarea name="descripcion"></textarea><label>Total cobrado / pactado</label><input id="monto_total" type="number" step="0.01" name="monto_total"><label>Abono que pagará</label><input type="number" step="0.01" name="abono"><label>Moneda</label><select id="moneda" name="moneda"><option>MXN</option><option>USD</option></select><label>Forma de pago</label><select name="forma_pago"><option>Efectivo</option><option>Tarjeta</option><option>Transferencia</option></select><label>Fecha de pago</label><input type="date" name="fecha_pago" value="{{hoy}}"><label>Próximo pago</label><input type="date" name="proximo_pago"><label>Comprobante</label><input type="file" name="comprobante"><button>Guardar y generar recibo PDF</button></form><script>function llenarExpediente(sel){var o=sel.options[sel.selectedIndex];document.getElementById('monto_total').value=o.dataset.total||'';if(o.dataset.moneda){document.getElementById('moneda').value=o.dataset.moneda;}}</script></div><div class="card print-list"><h2>Lista de pagos realizados <button class="no-print" onclick="window.print()">Imprimir lista</button></h2><table><tr><th>Folio</th><th>Cliente</th><th>Expediente</th><th>Concepto</th><th>Abono</th><th>Saldo</th><th>Moneda</th><th>Estatus</th><th>Próximo pago</th><th class="no-print">Recibo</th></tr>{% for p in datos %}<tr><td>{{p.folio}}</td><td>{{p.cliente.nombre if p.cliente else ""}}</td><td>{{p.expediente.numero if p.expediente else ""}}</td><td>{{p.concepto}}</td><td>{{p.abono}}</td><td>{{p.saldo}}</td><td>{{p.moneda}}</td><td>{{p.estatus}}</td><td>{{p.proximo_pago}}</td><td class="no-print"><a href="/recibo/{{p.id}}">Ver</a> | <a href="/recibo-pdf/{{p.id}}">PDF</a>{% if p.saldo <= 0 %} | <a href="/eliminar-cobranza/{{p.id}}" onclick="return confirm('¿Eliminar pago liquidado?')">Eliminar</a>{% endif %}</td></tr>{% endfor %}</table></div>""", clientes=Cliente.query.filter(Cliente.activo == True).order_by(Cliente.nombre).all(), expedientes=expedientes, datos=Cobranza.query.order_by(Cobranza.id.desc()).all(), hoy=local_date_iso())


@app.route("/eliminar-cobranza/<int:id>")
def eliminar_cobranza(id):
    if not req():
        return redirect("/login")
    p = Cobranza.query.get_or_404(id)
    if p.saldo > 0 and not admin():
        flash("Solo puedes eliminar pagos liquidados.")
        return redirect("/cobranza")
    db.session.delete(p)
    db.session.commit()
    flash("Registro de cobranza eliminado.")
    return redirect("/cobranza")


def recibo_html(p):
    return render_template_string("""<!doctype html><html><head><meta charset="utf-8"><title>Recibo {{p.folio}}</title><style>body{font-family:Arial;margin:35px;color:#111827}.recibo{border:2px solid #111827;padding:25px;border-radius:12px}.head{display:flex;align-items:center;border-bottom:3px solid #B38B2E;padding-bottom:15px}.logo{width:90px;height:90px;border-radius:50%;object-fit:cover;margin-right:18px}.grid{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:18px}.box{border:1px solid #ddd;padding:12px;border-radius:8px}.firma{margin-top:60px;border-top:1px solid #111;width:320px;text-align:center;padding-top:8px}.no-print{margin-bottom:20px}@media print{.no-print{display:none}}</style></head><body><div class="no-print"><button onclick="window.print()">Imprimir</button> <a href="/recibo-pdf/{{p.id}}">Descargar PDF</a></div><div class="recibo"><div class="head"><img class="logo" src="/logo"><div><h1>ALESI GRUPO JURÍDICO</h1><p>Recibo de pago</p></div></div><h2>Folio: {{p.folio}}</h2><div class="grid"><div class="box"><b>Fecha:</b> {{p.fecha_pago or p.creado_en.strftime('%Y-%m-%d')}}</div><div class="box"><b>Cliente:</b> {{p.cliente.nombre if p.cliente else ''}}</div><div class="box"><b>Expediente:</b> {{p.expediente.numero if p.expediente else ''}}</div><div class="box"><b>Concepto:</b> {{p.concepto}}</div><div class="box"><b>Monto total:</b> {{p.moneda}} {{'%.2f'|format(p.monto_total or 0)}}</div><div class="box"><b>Abono recibido:</b> {{p.moneda}} {{'%.2f'|format(p.abono or 0)}}</div><div class="box"><b>Saldo pendiente:</b> {{p.moneda}} {{'%.2f'|format(p.saldo or 0)}}</div><div class="box"><b>Forma de pago:</b> {{p.forma_pago}}</div><div class="box"><b>Próximo pago:</b> {{p.proximo_pago}}</div><div class="box"><b>Estatus:</b> {{p.estatus}}</div></div><p><b>Descripción / observaciones:</b><br>{{p.descripcion}}</p><p><b>Recibió:</b> {{p.usuario.nombre if p.usuario else ''}}</p><div class="firma">Firma de recibido</div></div></body></html>""", p=p)


@app.route("/recibo/<int:id>")
def recibo(id):
    if not req():
        return redirect("/login")
    p = Cobranza.query.get_or_404(id)
    return recibo_html(p)


def crear_recibo_pdf(p):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("AlesiTitle", parent=styles["Title"], textColor=colors.HexColor("#111827"), fontSize=20, spaceAfter=8)
    subtitle_style = ParagraphStyle("AlesiSub", parent=styles["Normal"], textColor=colors.HexColor("#B38B2E"), fontSize=12, spaceAfter=14)
    story = [Paragraph("ALESI GRUPO JURÍDICO", title_style), Paragraph("Recibo de pago", subtitle_style), Spacer(1, 10)]
    data = [
        ["Folio", p.folio or ""], ["Fecha", p.fecha_pago or p.creado_en.strftime("%Y-%m-%d")],
        ["Cliente", p.cliente.nombre if p.cliente else ""], ["Expediente", p.expediente.numero if p.expediente else ""],
        ["Concepto", p.concepto or ""], ["Monto total", f"{p.moneda or ''} {p.monto_total or 0:,.2f}"],
        ["Abono recibido", f"{p.moneda or ''} {p.abono or 0:,.2f}"], ["Saldo pendiente", f"{p.moneda or ''} {p.saldo or 0:,.2f}"],
        ["Forma de pago", p.forma_pago or ""], ["Próximo pago", p.proximo_pago or ""], ["Estatus", p.estatus or ""],
        ["Recibió", p.usuario.nombre if p.usuario else ""],
    ]
    table = Table(data, colWidths=[1.8*inch, 4.8*inch])
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(0,-1),colors.HexColor("#111827")),("TEXTCOLOR",(0,0),(0,-1),colors.white),("BACKGROUND",(1,0),(1,-1),colors.whitesmoke),("BOX",(0,0),(-1,-1),1,colors.HexColor("#B38B2E")),("INNERGRID",(0,0),(-1,-1),0.5,colors.lightgrey),("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),("VALIGN",(0,0),(-1,-1),"TOP"),("PADDING",(0,0),(-1,-1),8)]))
    story.append(table)
    story.append(Spacer(1, 18))
    story.append(Paragraph("<b>Descripción / observaciones:</b>", styles["Normal"]))
    story.append(Paragraph((p.descripcion or "").replace("\\n","<br/>"), styles["Normal"]))
    story.append(Spacer(1, 60))
    firma = Table([["Firma de recibido"]], colWidths=[3.5*inch])
    firma.setStyle(TableStyle([("LINEABOVE",(0,0),(-1,0),1,colors.black),("ALIGN",(0,0),(-1,-1),"CENTER"),("TOPPADDING",(0,0),(-1,-1),8)]))
    story.append(firma)
    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route("/recibo-pdf/<int:id>")
def recibo_pdf(id):
    if not req():
        return redirect("/login")
    p = Cobranza.query.get_or_404(id)
    return send_file(crear_recibo_pdf(p), mimetype="application/pdf", as_attachment=True, download_name=f"recibo_{p.folio or p.id}.pdf")


@app.route("/planeador")
def planeador():
    if not req():
        return redirect("/login")
    vista = request.args.get("vista", "mes")
    base_fecha = request.args.get("fecha") or local_date_iso()
    try:
        base = datetime.strptime(base_fecha, "%Y-%m-%d").date()
    except Exception:
        base = local_now().date()
    if vista == "dia":
        inicio = fin = base
    elif vista == "semana":
        inicio = base - timedelta(days=base.weekday())
        fin = inicio + timedelta(days=6)
    else:
        inicio = base.replace(day=1)
        if inicio.month == 12:
            fin = inicio.replace(year=inicio.year+1, month=1, day=1) - timedelta(days=1)
        else:
            fin = inicio.replace(month=inicio.month+1, day=1) - timedelta(days=1)
    eventos = []
    dias_semana = ["Lunes","Martes","Miércoles","Jueves","Viernes","Sábado","Domingo"]
    def en_rango(f):
        try:
            d = datetime.strptime(f, "%Y-%m-%d").date()
            return inicio <= d <= fin
        except Exception:
            return False
    for a in Audiencia.query.all():
        if a.fecha and en_rango(a.fecha):
            eventos.append({"fecha": a.fecha, "hora": a.hora or "", "tipo": "Audiencia", "texto": f"{a.expediente.numero if a.expediente else ''} - {a.titulo}"})
    for v in Movimiento.query.filter(Movimiento.fecha_limite != "").all():
        if v.fecha_limite and en_rango(v.fecha_limite):
            eventos.append({"fecha": v.fecha_limite, "hora": v.hora_limite or "", "tipo": "Vencimiento", "texto": f"{v.expediente.numero if v.expediente else ''} - {v.proxima_accion or v.titulo}"})
    for p in Cobranza.query.filter(Cobranza.proximo_pago != "", Cobranza.saldo > 0).all():
        if p.proximo_pago and en_rango(p.proximo_pago):
            eventos.append({"fecha": p.proximo_pago, "hora": "", "tipo": "Pago pendiente", "texto": f"{p.cliente.nombre if p.cliente else ''} - {p.moneda} {p.saldo}"})
    for c in CitaProspecto.query.filter(CitaProspecto.fecha >= local_date_iso()).all():
        if c.fecha and en_rango(c.fecha):
            eventos.append({"fecha": c.fecha, "hora": c.hora or "", "tipo": "Cita", "texto": f"{c.nombre} - {c.tipo_asunto} - Doc: {c.documentacion_solicitada or ''}"})
    eventos.sort(key=lambda x: (x["fecha"], x["hora"]))
    agrupado = {}
    for ev in eventos:
        agrupado.setdefault(ev["fecha"], []).append(ev)
    dias = {f: dias_semana[datetime.strptime(f, "%Y-%m-%d").weekday()] if f else "" for f in agrupado}
    return render("""<div class="card no-print"><h2>Planeador jurídico</h2><form method="get"><label>Vista</label><select name="vista"><option value="dia" {% if vista=='dia' %}selected{% endif %}>Día</option><option value="semana" {% if vista=='semana' %}selected{% endif %}>Semana</option><option value="mes" {% if vista=='mes' %}selected{% endif %}>Mes</option></select><label>Fecha base</label><input type="date" name="fecha" value="{{base_fecha}}"><button>Ver</button> <button type="button" onclick="window.print()">Imprimir calendario</button></form><p>Del {{inicio}} al {{fin}}.</p></div><div class="card solo-imprimir"><h2>Calendario ALESI Grupo Jurídico</h2><table><tr><th style="width:170px">Día</th><th>Actividades</th></tr>{% for fecha, lista in agrupado.items() %}<tr><td><b>{{dias[fecha]}}</b><br>{{fecha}}</td><td>{% for ev in lista %}<div style="margin-bottom:8px"><b>{{ev.tipo}}</b> {% if ev.hora %}{{ev.hora}}{% endif %}<br>{{ev.texto}}</div>{% endfor %}</td></tr>{% endfor %}</table></div>""", agrupado=agrupado, dias=dias, vista=vista, base_fecha=base_fecha, inicio=inicio, fin=fin)


@app.route("/notificaciones-internas", methods=["GET", "POST"])
def notificaciones_internas():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        titulo = request.form["titulo"]
        mensaje = request.form["mensaje"]
        expediente_id = request.form.get("expediente_id") or None
        grupo = "AVISO-" + local_now().strftime("%Y%m%d%H%M%S%f")
        enlace = f"/expediente/{expediente_id}" if expediente_id else ""
        for u in Usuario.query.filter_by(activo=True).all():
            notificar_interna(u.id, titulo, mensaje, enlace, creador_id=actual().id, grupo=grupo)
        if expediente_id:
            db.session.add(AvisoExpediente(expediente_id=expediente_id, usuario_id=actual().id, titulo=titulo, mensaje=mensaje, fecha_aviso=local_date_iso()))
            registrar_bitacora(int(expediente_id), "Aviso agregado", titulo + " - " + mensaje[:150])
        db.session.commit()
        flash("Aviso publicado para todos los usuarios.")
        return redirect("/notificaciones-internas")

    avisos = NotificacionInterna.query.filter_by(usuario_id=actual().id).order_by(NotificacionInterna.id.desc()).all()
    for aviso in avisos:
        aviso.leida = True
    db.session.commit()

    grupos = []
    creados = {}
    for a in NotificacionInterna.query.filter_by(creador_id=actual().id).order_by(NotificacionInterna.id.desc()).all():
        if not a.grupo:
            continue
        creados.setdefault(a.grupo, []).append(a)
    for grupo, lista in creados.items():
        grupos.append({"grupo": grupo, "titulo": lista[0].titulo, "total": len(lista), "leidos": sum(1 for x in lista if x.leida)})

    return render("""
    <div class="grid2">
      <div class="card">
        <h2>Agregar aviso interno</h2>
        <p class="small">El aviso se publicará para todos los usuarios del despacho.</p>
        <form method="post">
          <label>Título</label><input name="titulo" required>
          <label>Mensaje</label><textarea name="mensaje" required></textarea>
          <label>Relacionar con expediente</label>
          <select name="expediente_id"><option value="">Sin expediente</option>{% for e in expedientes %}<option value="{{e.id}}">{{e.numero}}</option>{% endfor %}</select>
          <button>Publicar aviso</button>
        </form>
      </div>
      <div class="card">
        <h2>Avisos internos</h2>
        <p class="small">Al abrir esta pestaña se marcan automáticamente como leídos y ya no aparecen en Inicio.</p>
        {% for a in avisos %}
        <div class="chat"><b>{{a.titulo}}</b> <span class="small">{{a.creado_en.strftime('%d/%m/%Y %H:%M')}}</span><br>{{a.mensaje}}<br>{% if a.enlace %}<a class="btn btn2" href="{{a.enlace}}">Abrir</a>{% endif %} <a class="btn btnDark" href="/eliminar-aviso-interno/{{a.id}}" onclick="return confirm('¿Eliminar este aviso de tu lista?')">Eliminar de mi lista</a></div>
        {% endfor %}
      </div>
    </div>
    <div class="card">
      <h2>Avisos que publiqué</h2>
      <table><tr><th>Aviso</th><th>Leídos</th><th>Acción</th></tr>{% for g in grupos %}<tr><td>{{g.titulo}}</td><td>{{g.leidos}} / {{g.total}}</td><td>{% if g.leidos == g.total %}<a class="btn btnDark" href="/eliminar-aviso-grupo/{{g.grupo}}" onclick="return confirm('¿Eliminar este aviso para todos?')">Eliminar para todos</a>{% else %}<span class="small">Se podrá eliminar cuando todos lo lean</span>{% endif %}</td></tr>{% endfor %}</table>
    </div>
    """, avisos=avisos, grupos=grupos, expedientes=Expediente.query.filter(Expediente.estado != "Eliminado").order_by(Expediente.numero).all())


@app.route("/eliminar-aviso-interno/<int:id>")
def eliminar_aviso_interno(id):
    if not req():
        return redirect("/login")
    aviso = NotificacionInterna.query.get_or_404(id)
    if aviso.usuario_id != actual().id and not admin():
        flash("No tienes permiso para eliminar este aviso.")
        return redirect("/notificaciones-internas")
    db.session.delete(aviso)
    db.session.commit()
    flash("Aviso eliminado.")
    return redirect("/notificaciones-internas")



@app.route("/eliminar-aviso-grupo/<grupo>")
def eliminar_aviso_grupo(grupo):
    if not req():
        return redirect("/login")
    avisos = NotificacionInterna.query.filter_by(grupo=grupo).all()
    if not avisos:
        flash("Aviso no encontrado.")
        return redirect("/notificaciones-internas")
    if not (admin() or avisos[0].creador_id == actual().id):
        flash("No tienes permiso para eliminar este aviso.")
        return redirect("/notificaciones-internas")
    if not admin() and not all(a.leida for a in avisos):
        flash("Aún no todos los usuarios han leído este aviso.")
        return redirect("/notificaciones-internas")
    for a in avisos:
        db.session.delete(a)
    db.session.commit()
    flash("Aviso eliminado para todos.")
    return redirect("/notificaciones-internas")

@app.route("/avisos-marcar-leidos")
def avisos_marcar_leidos():
    if not req():
        return redirect("/login")
    for aviso in NotificacionInterna.query.filter_by(usuario_id=actual().id).all():
        aviso.leida = True
    db.session.commit()
    flash("Avisos marcados como leídos.")
    return redirect("/notificaciones-internas")


@app.route("/jurisprudencia")
def jurisprudencia():
    if not req():
        return redirect("/login")
    q = request.args.get("q", "").strip()
    enlace = ""
    if q:
        enlace = "https://www.google.com/search?q=" + quote_plus(q + " jurisprudencia Suprema Corte Semanario Judicial site:scjn.gob.mx OR site:sjf2.scjn.gob.mx")
    return render("""<div class="card"><h2>Jurisprudencia</h2><form><label>Concepto jurídico</label><input name="q" value="{{q}}" placeholder="Ejemplo: pensión alimenticia, despido injustificado"><button>Preparar búsqueda ampliada</button></form>{% if q %}<br><a class="btn btn2" target="_blank" href="{{enlace}}">Abrir búsqueda ampliada</a>{% endif %}</div>""", q=q, enlace=enlace)


@app.route("/boletin-judicial")
def boletin_judicial():
    if not req():
        return redirect("/login")
    return render("""<div class="card"><h2>Boletín Judicial de Baja California</h2><p>Acceso directo al Boletín Judicial.</p><a class="btn btn2" target="_blank" href="https://www.pjbc.gob.mx/boletin_judicial.aspx">Abrir Boletín Judicial BC</a></div>""")


@app.route("/teja")
def teja():
    if not req():
        return redirect("/login")
    return render("""<div class="card"><h2>TEJA Baja California</h2><p>Accesos directos al Tribunal Estatal de Justicia Administrativa de Baja California.</p><a class="btn" target="_blank" href="https://tejabc.mx/">Abrir TEJA BC</a> <a class="btn btn2" target="_blank" href="https://tejabc.mx/buscarlistas">Buscar listas TEJA</a></div>""")


@app.route("/citas", methods=["GET", "POST"])
def citas():
    if not req():
        return redirect("/login")
    if request.method == "POST":
        cita = CitaProspecto(
            nombre=request.form.get("nombre", ""),
            telefono=request.form.get("telefono", ""),
            correo=request.form.get("correo", ""),
            materia=request.form.get("materia", ""),
            tipo_asunto=request.form.get("tipo_asunto", ""),
            fecha=request.form.get("fecha", ""),
            hora=request.form.get("hora", ""),
            documentacion_solicitada=request.form.get("documentacion_solicitada", ""),
            estatus="Agendada",
            observaciones=request.form.get("observaciones", ""),
        )
        db.session.add(cita)
        db.session.commit()
        flash("Cita registrada.")
        return redirect("/citas")
    datos = CitaProspecto.query.filter(CitaProspecto.fecha >= local_date_iso()).order_by(CitaProspecto.fecha, CitaProspecto.hora).all()
    return render("""<div class="grid2"><div class="card no-print"><h2>Registrar cita</h2><form method="post"><label>Nuevo cliente</label><input name="nombre" required><label>Teléfono</label><input name="telefono"><label>Correo</label><input name="correo"><label>Materia</label><select name="materia">{% for m in materias %}<option>{{m}}</option>{% endfor %}</select><label>Tipo de asunto</label><select name="tipo_asunto">{% for t in tipos %}<option>{{t.nombre}}</option>{% endfor %}</select><label>Fecha de cita</label><input type="date" name="fecha" required><label>Hora</label><input type="time" name="hora"><label>Documentación solicitada</label><textarea name="documentacion_solicitada" placeholder="Ej. acta de matrimonio, INE, comprobantes, etc."></textarea><label>Observaciones</label><textarea name="observaciones"></textarea><button>Guardar cita</button></form></div><div class="card print-list"><h2>Citas registradas <button class="no-print" onclick="window.print()">Imprimir lista</button></h2><p class="small">Solo se muestran citas vigentes. Las citas pasadas se ocultan automáticamente del planeador y de esta lista.</p><table><tr><th>Fecha</th><th>Hora</th><th>Cliente</th><th>Materia</th><th>Tipo asunto</th><th>Documentación solicitada</th><th class="no-print">Acción</th></tr>{% for c in datos %}<tr><td>{{c.fecha}}</td><td>{{c.hora}}</td><td>{{c.nombre}}</td><td>{{c.materia}}</td><td>{{c.tipo_asunto}}</td><td>{{c.documentacion_solicitada}}</td><td class="no-print"><a class="btn btn2" href="/cita/{{c.id}}">Editar</a> <a class="btn btnDark" href="/eliminar-cita/{{c.id}}" onclick="return confirm('¿Eliminar cita?')">Eliminar</a></td></tr>{% endfor %}</table></div></div>""", datos=datos, materias=MATERIAS, tipos=TipoAsunto.query.order_by(TipoAsunto.materia, TipoAsunto.nombre).all())


@app.route("/cita/<int:id>", methods=["GET", "POST"])
def cita(id):
    if not req():
        return redirect("/login")
    c = CitaProspecto.query.get_or_404(id)
    if request.method == "POST":
        c.nombre = request.form.get("nombre", c.nombre)
        c.telefono = request.form.get("telefono", "")
        c.correo = request.form.get("correo", "")
        c.materia = request.form.get("materia", "")
        c.tipo_asunto = request.form.get("tipo_asunto", "")
        c.fecha = request.form.get("fecha", "")
        c.hora = request.form.get("hora", "")
        c.documentacion_solicitada = request.form.get("documentacion_solicitada", "")
        c.estatus = request.form.get("estatus", "Agendada")
        c.observaciones = request.form.get("observaciones", "")
        db.session.commit()
        flash("Cita actualizada.")
        return redirect("/citas")
    return render("""<div class="card"><h2>Editar cita</h2><form method="post"><label>Cliente</label><input name="nombre" value="{{c.nombre}}" required><label>Teléfono</label><input name="telefono" value="{{c.telefono or ''}}"><label>Correo</label><input name="correo" value="{{c.correo or ''}}"><label>Materia</label><select name="materia">{% for m in materias %}<option {% if c.materia==m %}selected{% endif %}>{{m}}</option>{% endfor %}</select><label>Tipo de asunto</label><input name="tipo_asunto" value=""><label>Fecha</label><input type="date" name="fecha" value="{{c.fecha or ''}}"><label>Hora</label><input type="time" name="hora" value="{{c.hora or ''}}"><label>Documentación solicitada</label><textarea name="documentacion_solicitada">{{c.documentacion_solicitada or ''}}</textarea><label>Estatus</label><select name="estatus"><option>Agendada</option><option>Atendida</option><option>Reprogramada</option><option>Cancelada</option></select><label>Observaciones</label><textarea name="observaciones">{{c.observaciones or ''}}</textarea><button>Guardar cambios</button></form></div>""", c=c, materias=MATERIAS)


@app.route("/eliminar-cita/<int:id>")
def eliminar_cita(id):
    if not req():
        return redirect("/login")
    c = CitaProspecto.query.get_or_404(id)
    db.session.delete(c)
    db.session.commit()
    flash("Cita eliminada.")
    return redirect("/citas")


UNIDADES = ["", "UNO", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
ESPECIALES = {
    10: "DIEZ", 11: "ONCE", 12: "DOCE", 13: "TRECE", 14: "CATORCE", 15: "QUINCE",
    16: "DIECISÉIS", 17: "DIECISIETE", 18: "DIECIOCHO", 19: "DIECINUEVE",
    20: "VEINTE", 21: "VEINTIUNO", 22: "VEINTIDÓS", 23: "VEINTITRÉS", 24: "VEINTICUATRO",
    25: "VEINTICINCO", 26: "VEINTISÉIS", 27: "VEINTISIETE", 28: "VEINTIOCHO", 29: "VEINTINUEVE"
}
DECENAS = {30: "TREINTA", 40: "CUARENTA", 50: "CINCUENTA", 60: "SESENTA", 70: "SETENTA", 80: "OCHENTA", 90: "NOVENTA"}
CENTENAS = {100: "CIEN", 200: "DOSCIENTOS", 300: "TRESCIENTOS", 400: "CUATROCIENTOS", 500: "QUINIENTOS", 600: "SEISCIENTOS", 700: "SETECIENTOS", 800: "OCHOCIENTOS", 900: "NOVECIENTOS"}
MESES = ["", "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO", "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]


def numero_entero_letra(n):
    n = int(n)
    if n == 0:
        return "CERO"
    if n < 10:
        return UNIDADES[n]
    if n < 30:
        return ESPECIALES[n]
    if n < 100:
        d = (n // 10) * 10
        r = n % 10
        return DECENAS[d] if r == 0 else DECENAS[d] + " Y " + UNIDADES[r]
    if n < 1000:
        if n in CENTENAS:
            return CENTENAS[n]
        c = (n // 100) * 100
        r = n % 100
        return ("CIENTO" if c == 100 else CENTENAS[c]) + " " + numero_entero_letra(r)
    if n < 1000000:
        miles = n // 1000
        r = n % 1000
        texto = "MIL" if miles == 1 else numero_entero_letra(miles) + " MIL"
        return texto if r == 0 else texto + " " + numero_entero_letra(r)
    millones = n // 1000000
    r = n % 1000000
    texto = "UN MILLÓN" if millones == 1 else numero_entero_letra(millones) + " MILLONES"
    return texto if r == 0 else texto + " " + numero_entero_letra(r)


def cantidad_en_letra(valor):
    try:
        numero = float(str(valor).replace(",", ""))
    except Exception:
        numero = 0
    entero = int(numero)
    centavos = int(round((numero - entero) * 100))
    return f"{numero_entero_letra(entero)} PESOS {centavos:02d}/100 MONEDA NACIONAL"


def fecha_contrato_texto(fecha_iso):
    try:
        f = datetime.strptime(fecha_iso, "%Y-%m-%d")
        return str(f.day), MESES[f.month], str(f.year), f"{f.day} DE {MESES[f.month]} DEL {f.year}"
    except Exception:
        hoy = local_now()
        return str(hoy.day), MESES[hoy.month], str(hoy.year), f"{hoy.day} DE {MESES[hoy.month]} DEL {hoy.year}"


def ptxt(texto, style):
    return Paragraph(str(texto or "").replace("&", "&amp;").replace("\n", "<br/>"), style)


def encabezado_pie(canvas, doc):
    canvas.saveState()
    w, h = letter
    logo_path = "logo_alesi.png"
    if os.path.exists(logo_path):
        try:
            canvas.drawImage(logo_path, 40, h - 90, width=95, height=70, preserveAspectRatio=True, mask='auto')
        except Exception:
            pass
    canvas.setStrokeColor(colors.HexColor("#B38B2E"))
    canvas.setLineWidth(1)
    canvas.line(38, 52, w - 38, 52)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#6b7280"))
    canvas.drawCentredString(w/2, 38, "ALESI GRUPO JURÍDICO | Cuauhtémoc 5048, Col. Ampliación Gabriel Rodríguez, 22207, Tijuana, Baja California | 664 565 7298 | lic.rosavazquezm@gmail.com")
    canvas.drawRightString(w - 38, 24, f"Página {doc.page}")
    canvas.restoreState()



def contrato_plantilla_path():
    posibles = [
        "contrato_alesi.docx",
        "1.-CONTRATO DE PRESTACION DE SERVICIOS PROFESIONALES (2)(2).docx",
        "1.-CONTRATO DE PRESTACION DE SERVICIOS PROFESIONALES (2).docx",
    ]
    for nombre in posibles:
        if os.path.exists(nombre):
            return nombre
    return None


def limpiar_prestador(nombre):
    nombre = (nombre or "ROSA ISELA VÁZQUEZ MEDINA").strip()
    nombre = nombre.replace("LIC. ", "").replace("LIC.", "").strip()
    return nombre.upper()


def set_paragraph_text(paragraph, nuevo_texto):
    if paragraph.runs:
        paragraph.runs[0].text = str(nuevo_texto)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(nuevo_texto))


def fill_first_blank(paragraph, valor):
    texto = paragraph.text
    nuevo = re.sub(r"_{3,}(?:\s+_{3,})*", str(valor or ""), texto, count=1)
    set_paragraph_text(paragraph, nuevo)


def generar_contrato_docx_desde_plantilla(cliente, expediente, datos):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Falta instalar python-docx. Agrega python-docx==1.1.2 en requirements.txt") from exc

    plantilla = contrato_plantilla_path()
    if not plantilla:
        raise FileNotFoundError("No se encontró contrato_alesi.docx en la raíz del proyecto.")

    fecha_iso = datos.get("fecha_contrato") or local_date_iso()
    dia, mes, anio, fecha_larga = fecha_contrato_texto(fecha_iso)

    cliente_nombre = (datos.get("cliente_nombre") or cliente.nombre or "").upper()
    cliente_tel = datos.get("telefono") or cliente.telefono or ""
    cliente_correo = datos.get("correo") or cliente.correo or ""
    domicilio = datos.get("domicilio") or cliente.direccion or ""
    curp = datos.get("curp") or cliente.curp or ""
    rfc = datos.get("rfc") or cliente.rfc or ""
    fecha_nacimiento = datos.get("fecha_nacimiento") or ""
    estado_civil = datos.get("estado_civil") or ""
    ocupacion = datos.get("ocupacion") or ""
    identificacion = datos.get("identificacion") or ""
    numero_identificacion = datos.get("numero_identificacion") or ""
    asunto = datos.get("tipo_asunto") or (expediente.tipo_asunto if expediente else "") or ""
    monto_total = float(datos.get("monto_total") or 0)
    monto_inicio = float(datos.get("monto_inicio") or 0)
    monto_abono = float(datos.get("monto_abono") or 0)
    periodicidad = datos.get("periodicidad_abono") or ""
    prestador = limpiar_prestador(datos.get("prestador"))
    prestador_correo = datos.get("prestador_correo") or "lic.rosavazquezm@gmail.com"
    prestador_telefono = datos.get("prestador_telefono") or "6645657298"

    registro = ContratoCliente(
        cliente_id=cliente.id,
        expediente_id=expediente.id if expediente else None,
        usuario_id=actual().id if actual() else None,
        cliente_nombre=cliente_nombre,
        tipo_asunto=asunto,
        monto_total=monto_total,
        monto_letra=cantidad_en_letra(monto_total),
        forma_pago=datos.get("forma_pago") or "",
        fecha_contrato=fecha_larga,
        personal_atendio=prestador,
        telefono_cliente=cliente_tel,
        correo_cliente=cliente_correo,
    )
    db.session.add(registro)
    db.session.commit()

    doc = Document(plantilla)

    for p in doc.paragraphs:
        txt = p.text

        if "Tijuana, Baja California, a ____ de __________________ del ______." in txt:
            set_paragraph_text(p, f"Tijuana, Baja California, a {dia} de {mes} del {anio}.")
            continue

        if "LICENCIADA EN DERECHO" in txt and "EL CLIENTE" in txt:
            set_paragraph_text(
                p,
                f'CONTRATO DE PRESTACIÓN DE SERVICIOS JURÍDICOS PROFESIONALES QUE CELEBRAN POR UNA PARTE EL (LA) LICENCIADA EN DERECHO {prestador}, A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "EL PRESTADOR"; Y POR LA OTRA PARTE EL(LA) C. {cliente_nombre}, A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "EL CLIENTE", QUIENES MANIFIESTAN SU VOLUNTAD DE OBLIGARSE AL TENOR DE LAS SIGUIENTES:'
            )
            continue

        if txt.startswith("Correo electrónico:"):
            set_paragraph_text(p, "Correo electrónico: " + prestador_correo)
            continue

        if txt.startswith("Teléfono:") and "Identificación" not in txt:
            # Primer teléfono dentro de declaraciones del prestador y luego cliente. Se distingue por el texto actual.
            if "______________________________________________" in txt:
                set_paragraph_text(p, "Teléfono: " + prestador_telefono)
                continue

        if "d) Que señala como domicilio" in txt:
            set_paragraph_text(p, "d) Que señala como domicilio para oír y recibir notificaciones el ubicado en: " + domicilio)
            continue

        if txt.startswith("Fecha de Nacimiento:"):
            set_paragraph_text(p, "Fecha de Nacimiento: " + fecha_nacimiento)
            continue
        if txt.startswith("CURP:"):
            set_paragraph_text(p, "CURP: " + curp)
            continue
        if txt.startswith("RFC:"):
            set_paragraph_text(p, "RFC: " + rfc)
            continue
        if txt.startswith("Estado Civil:"):
            set_paragraph_text(p, "Estado Civil: " + estado_civil)
            continue
        if txt.startswith("Ocupación:"):
            set_paragraph_text(p, "Ocupación: " + ocupacion)
            continue
        if txt.startswith("Teléfono:") and "____" in txt:
            set_paragraph_text(p, "Teléfono: " + cliente_tel)
            continue
        if txt.startswith("Correo Electrónico:"):
            set_paragraph_text(p, "Correo Electrónico: " + cliente_correo)
            continue
        if txt.startswith("Identificación Oficial:"):
            set_paragraph_text(p, "Identificación Oficial: " + identificacion)
            continue
        if txt.startswith("Número de Identificación:"):
            set_paragraph_text(p, "Número de Identificación: " + numero_identificacion)
            continue

        if "respecto del siguiente asunto:" in txt:
            set_paragraph_text(
                p,
                "EL PRESTADOR se obliga a proporcionar a EL CLIENTE servicios jurídicos profesionales consistentes en asesoría, orientación legal, elaboración de documentos jurídicos, representación, gestión administrativa, negociación, mediación, conciliación, seguimiento procesal y, en su caso, litigio ante autoridades administrativas, judiciales o de cualquier otra naturaleza, respecto del siguiente asunto: " + asunto
            )
            continue

        if "por concepto de HONORARIOS" in txt:
            set_paragraph_text(
                p,
                f"Por los servicios profesionales objeto del presente contrato, EL CLIENTE se obliga a pagar a EL PRESTADOR la cantidad de: ${monto_total:,.2f} PESOS M.N ({cantidad_en_letra(monto_total)}) por concepto de HONORARIOS"
            )
            continue

        if "por concepto DE INICIO" in txt:
            set_paragraph_text(
                p,
                f"${monto_inicio:,.2f} PESOS M.N ({cantidad_en_letra(monto_inicio)}) por concepto DE INICIO, pagaderos a la firma del presente contrato."
            )
            continue

        if "POR CONCEPTO DE ABONO" in txt and "PAGADEROS" in txt:
            set_paragraph_text(
                p,
                f"${monto_abono:,.2f} PESOS M.N ({cantidad_en_letra(monto_abono)}) POR CONCEPTO DE ABONO PAGADEROS {periodicidad} y hasta la liquidación."
            )
            continue

        if "Leído que fue el presente contrato" in txt and "____ de __________________ del ______" in txt:
            set_paragraph_text(
                p,
                f"Leído que fue el presente contrato y enteradas las partes de su contenido, alcance y consecuencias legales, lo firman por duplicado en la ciudad de Tijuana, Baja California, a {dia} de {mes} del {anio}."
            )
            continue

        if txt.strip().startswith("LIC.") and "Nombre:" in txt:
            set_paragraph_text(p, f"{prestador}                                                                       Nombre: {cliente_nombre}")
            continue

        if txt.startswith("Teléfono:") and "________________________" in txt:
            set_paragraph_text(p, "Teléfono: " + cliente_tel)
            continue

        if txt.startswith("Identificación Oficial No."):
            set_paragraph_text(p, "Identificación Oficial No. " + numero_identificacion)
            continue

    salida = BytesIO()
    doc.save(salida)
    salida.seek(0)

    nombre_base = "contrato_" + cliente_nombre.replace(" ", "_").replace("/", "_")
    return salida, nombre_base


def convertir_docx_bytes_a_pdf(docx_bytes, nombre_base):
    tmpdir = tempfile.mkdtemp(prefix="alesi_contrato_")
    try:
        docx_path = os.path.join(tmpdir, nombre_base + ".docx")
        pdf_path = os.path.join(tmpdir, nombre_base + ".pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes.getvalue())

        soffice = shutil.which("libreoffice") or shutil.which("soffice")
        if not soffice:
            raise RuntimeError("No está instalado LibreOffice. Sube Aptfile con libreoffice para convertir a PDF en Render.")

        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
        )

        if not os.path.exists(pdf_path):
            encontrados = [x for x in os.listdir(tmpdir) if x.lower().endswith(".pdf")]
            if encontrados:
                pdf_path = os.path.join(tmpdir, encontrados[0])

        with open(pdf_path, "rb") as f:
            pdf_data = f.read()

        return BytesIO(pdf_data)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def contrato_pdf_response(cliente, expediente, datos):
    docx_bytes, nombre_base = generar_contrato_docx_desde_plantilla(cliente, expediente, datos)
    pdf_bytes = convertir_docx_bytes_a_pdf(docx_bytes, nombre_base)
    pdf_bytes.seek(0)
    return send_file(pdf_bytes, as_attachment=True, download_name=nombre_base + ".pdf", mimetype="application/pdf")


def contrato_word_response(cliente, expediente, datos):
    docx_bytes, nombre_base = generar_contrato_docx_desde_plantilla(cliente, expediente, datos)
    docx_bytes.seek(0)
    return send_file(docx_bytes, as_attachment=True, download_name=nombre_base + ".docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/monto-letra")
def monto_letra_api():
    return Response(cantidad_en_letra(request.args.get("monto", "0")), mimetype="text/plain; charset=utf-8")



@app.route("/contratos")
def contratos():
    if not req():
        return redirect("/login")
    clientes = Cliente.query.order_by(Cliente.nombre).all()
    contratos = ContratoCliente.query.order_by(ContratoCliente.id.desc()).limit(50).all()
    return render("""
    <div class="card">
        <h2>Contratos de prestación de servicios</h2>
        <p>Selecciona un cliente para llenar y descargar su contrato en PDF.</p>
        <table>
            <tr><th>Cliente</th><th>Teléfono</th><th>Correo</th><th>Acción</th></tr>
            {% for c in clientes %}
            <tr>
                <td>{{c.nombre}}</td>
                <td>{{c.telefono}}</td>
                <td>{{c.correo}}</td>
                <td><a class="btn btn2" href="/contrato-cliente/{{c.id}}">Crear contrato PDF</a></td>
            </tr>
            {% endfor %}
        </table>
    </div>
    <div class="card">
        <h2>Últimos contratos generados</h2>
        <table>
            <tr><th>Cliente</th><th>Asunto</th><th>Total</th><th>Fecha</th><th>Crear nuevamente</th></tr>
            {% for ct in contratos %}
            <tr>
                <td>{{ct.cliente.nombre if ct.cliente else ct.cliente_nombre}}</td>
                <td>{{ct.tipo_asunto}}</td>
                <td>{{ct.monto_total}}</td>
                <td>{{ct.fecha_contrato}}</td>
                <td>{% if ct.cliente %}<a class="btn" href="/contrato-cliente/{{ct.cliente.id}}">Abrir formulario</a>{% endif %}</td>
            </tr>
            {% endfor %}
        </table>
    </div>
    """, clientes=clientes, contratos=contratos)

@app.route("/contrato-cliente/<int:id>", methods=["GET", "POST"])
@app.route("/contrato/<int:id>", methods=["GET", "POST"])
@app.route("/generar-contrato/<int:id>", methods=["GET", "POST"])
@app.route("/contratos/cliente/<int:id>", methods=["GET", "POST"])
def contrato_cliente(id):
    if not req():
        return redirect("/login")
    c = Cliente.query.get_or_404(id)
    exps = Expediente.query.filter_by(cliente_id=id).filter(Expediente.estado != "Eliminado").order_by(Expediente.id.desc()).all()
    if request.method == "POST":
        expediente = db.session.get(Expediente, int(request.form.get("expediente_id"))) if request.form.get("expediente_id") else None
        return contrato_pdf_response(c, expediente, request.form)
    fecha_hoy = local_date_iso()
    return render("""
    <div class="card">
        <h2>Crear contrato de prestación de servicios jurídicos profesionales</h2>
        <p>Llena los datos en blanco y al guardar se descargará el contrato en PDF listo para imprimir.</p>
        <form method="post">
            <label>Cliente</label>
            <input name="cliente_nombre" value="{{c.nombre or ''}}" required>
            <label>Teléfono del cliente</label>
            <input name="telefono" value="{{c.telefono or ''}}">
            <label>Correo del cliente</label>
            <input name="correo" value="{{c.correo or ''}}">
            <label>Domicilio del cliente</label>
            <textarea name="domicilio">{{c.direccion or ''}}</textarea>
            <label>Fecha de nacimiento</label>
            <input type="date" name="fecha_nacimiento">
            <label>CURP</label>
            <input name="curp" value="{{c.curp or ''}}">
            <label>RFC</label>
            <input name="rfc" value="{{c.rfc or ''}}">
            <label>Estado civil</label>
            <input name="estado_civil">
            <label>Ocupación</label>
            <input name="ocupacion">
            <label>Identificación oficial</label>
            <input name="identificacion" placeholder="INE, pasaporte, licencia, etc.">
            <label>Número de identificación</label>
            <input name="numero_identificacion">
            <label>Expediente / asunto relacionado</label>
            <select name="expediente_id">
                <option value="">Usar datos generales del cliente</option>
                {% for e in exps %}<option value="{{e.id}}">{{e.numero}} - {{e.tipo_asunto}} - {{e.moneda_cobro}} {{e.cobro_total}}</option>{% endfor %}
            </select>
            <label>Tipo de asunto / servicio contratado</label>
            <input name="tipo_asunto" value="" required>
            <label>Fecha del contrato</label>
            <input type="date" name="fecha_contrato" value="{{fecha_hoy}}" required>
            <label>Total de honorarios</label>
            <input type="number" step="0.01" name="monto_total" id="monto_total" value="0" oninput="actualizarLetra()" required>
            <label>Cantidad en letra automática</label>
            <input id="monto_letra_preview" value="" readonly>
            <label>Pago inicial</label>
            <input type="number" step="0.01" name="monto_inicio" value="0">
            <label>Abono o parcialidad</label>
            <input type="number" step="0.01" name="monto_abono" value="0">
            <label>Periodicidad del abono</label>
            <input name="periodicidad_abono" placeholder="Ej. semanal, quincenal, mensual, en cada recuperación, etc.">
            <label>Forma de cobro / observación de pago</label>
            <textarea name="forma_pago" placeholder="Ej. pago inicial a la firma y abonos quincenales hasta liquidar; 50% de lo recuperado; etc."></textarea>
            <label>Licenciada / prestador</label>
            <input name="prestador" value="ROSA ISELA VÁZQUEZ MEDINA">
            <label>Personal que atendió y firma</label>
            <input name="personal_atendio" value="{{usuario.nombre if usuario else 'LIC. ROSA ISELA VÁZQUEZ MEDINA'}}">
            <button>Guardar y descargar contrato PDF</button>
        </form>
    </div>
    <script>
    function actualizarLetra(){
        const monto = document.getElementById('monto_total').value || '0';
        fetch('/monto-letra?monto=' + encodeURIComponent(monto))
          .then(r => r.text())
          .then(t => { document.getElementById('monto_letra_preview').value = t; })
          .catch(() => {});
    }
    window.onload = actualizarLetra;
    </script>
    """, c=c, exps=exps, fecha_hoy=fecha_hoy)


@app.route("/apariencia", methods=["GET", "POST"])
def apariencia():
    if not req() or not admin():
        return redirect("/")
    a = Apariencia.query.first()
    if not a:
        a = Apariencia()
        db.session.add(a)
        db.session.commit()
    if request.method == "POST":
        a.color_principal = request.form.get("color_principal", a.color_principal)
        a.color_secundario = request.form.get("color_secundario", a.color_secundario)
        a.color_dorado = request.form.get("color_dorado", a.color_dorado)
        a.fondo = request.form.get("fondo", a.fondo)
        a.fuente = request.form.get("fuente", a.fuente)
        logo = subir(request.files.get("logo"))
        if logo:
            a.logo_url = logo
        db.session.commit()
        flash("Apariencia actualizada.")
        return redirect("/apariencia")
    return render("""<div class="card"><h2>Apariencia del sistema</h2>{% if a.logo_url %}<p><img src="{{a.logo_url}}" style="max-width:160px;border:1px solid #ddd;border-radius:10px"></p>{% endif %}<form method="post" enctype="multipart/form-data"><label>Nuevo logo</label><input type="file" name="logo"><label>Color principal</label><input type="color" name="color_principal" value="{{a.color_principal}}"><label>Color secundario</label><input type="color" name="color_secundario" value="{{a.color_secundario}}"><label>Color dorado/acento</label><input type="color" name="color_dorado" value="{{a.color_dorado}}"><label>Color de fondo</label><input type="color" name="fondo" value="{{a.fondo}}"><label>Tipo de letra</label><select name="fuente"><option {% if a.fuente=='Arial' %}selected{% endif %}>Arial</option><option {% if a.fuente=='Calibri' %}selected{% endif %}>Calibri</option><option {% if a.fuente=='Georgia' %}selected{% endif %}>Georgia</option><option {% if a.fuente=='Times New Roman' %}selected{% endif %}>Times New Roman</option><option {% if a.fuente=='Verdana' %}selected{% endif %}>Verdana</option></select><button>Guardar cambios</button></form></div>""", a=a)


@app.route("/notificaciones")
def notificaciones():
    if not req() or not admin():
        return redirect("/")
    total = revisar_y_enviar_notificaciones()
    return render("""<div class="card"><h2>Notificaciones</h2><p>Procesadas: <b>{{total}}</b></p></div>""", total=total)


# =========================
# RESPALDOS MANUALES EN EXCEL
# =========================

@app.route("/respaldos")
def respaldos():
    if not req() or not admin():
        return redirect("/")

    return render("""
    <div class="card">
        <h2>Respaldos del Sistema</h2>
        <p>Descarga una copia completa de la información del sistema en Excel para guardarla en tu computadora.</p>
        <br>
        <a class="btn" href="/descargar-respaldo">Descargar respaldo completo en Excel</a>
    </div>
    """)


@app.route("/descargar-respaldo")
def descargar_respaldo():
    if not req() or not admin():
        return redirect("/")

    wb = Workbook()

    # CLIENTES
    ws = wb.active
    ws.title = "Clientes"
    ws.append(["ID", "Nombre", "Teléfono", "Correo", "RFC", "CURP", "Materia", "Tipo asunto", "Cobro total", "Moneda", "Dirección", "Observaciones", "Creado en"])
    for c in Cliente.query.order_by(Cliente.id).all():
        ws.append([
            c.id,
            c.nombre,
            c.telefono,
            c.correo,
            c.rfc,
            c.curp,
            c.materia,
            c.tipo_asunto,
            c.cobro_total,
            c.moneda_cobro,
            c.direccion,
            c.observaciones,
            c.creado_en.strftime("%d/%m/%Y %H:%M") if c.creado_en else "",
        ])

    # EXPEDIENTES
    ws = wb.create_sheet("Expedientes")
    ws.append(["ID", "Número", "Cliente", "Materia", "Tipo asunto", "Autoridad", "Actor", "Demandado", "Estado", "Prioridad", "Responsable", "Fecha inicio", "Observaciones", "Propietario", "Creado en"])
    for e in Expediente.query.order_by(Expediente.id).all():
        ws.append([
            e.id,
            e.numero,
            e.cliente.nombre if e.cliente else "",
            e.materia,
            e.tipo_asunto,
            e.autoridad,
            e.actor,
            e.demandado,
            e.estado,
            e.prioridad,
            e.responsable,
            e.fecha_inicio,
            e.observaciones,
            e.propietario.nombre if e.propietario else "",
            e.creado_en.strftime("%d/%m/%Y %H:%M") if e.creado_en else "",
        ])

    # COBRANZA
    ws = wb.create_sheet("Cobranza")
    ws.append(["ID", "Folio", "Cliente", "Expediente", "Usuario", "Concepto", "Descripción", "Monto total", "Abono", "Saldo", "Moneda", "Forma de pago", "Fecha pago", "Próximo pago", "Estatus", "Comprobante", "Creado en"])
    for p in Cobranza.query.order_by(Cobranza.id).all():
        ws.append([
            p.id,
            p.folio,
            p.cliente.nombre if p.cliente else "",
            p.expediente.numero if p.expediente else "",
            p.usuario.nombre if p.usuario else "",
            p.concepto,
            p.descripcion,
            p.monto_total,
            p.abono,
            p.saldo,
            p.moneda,
            p.forma_pago,
            p.fecha_pago,
            p.proximo_pago,
            p.estatus,
            p.comprobante_url,
            p.creado_en.strftime("%d/%m/%Y %H:%M") if p.creado_en else "",
        ])

    # AUDIENCIAS
    ws = wb.create_sheet("Audiencias")
    ws.append(["ID", "Expediente", "Usuario", "Título", "Fecha", "Hora", "Autoridad", "Sala", "Modalidad", "Enlace", "Observaciones", "Notificar", "Creado en"])
    for a in Audiencia.query.order_by(Audiencia.id).all():
        ws.append([
            a.id,
            a.expediente.numero if a.expediente else "",
            a.usuario.nombre if a.usuario else "",
            a.titulo,
            a.fecha,
            a.hora,
            a.autoridad,
            a.sala,
            a.modalidad,
            a.enlace,
            a.observaciones,
            "Sí" if a.notificar else "No",
            a.creado_en.strftime("%d/%m/%Y %H:%M") if a.creado_en else "",
        ])

    # VENCIMIENTOS / MOVIMIENTOS
    ws = wb.create_sheet("Vencimientos")
    ws.append(["ID", "Expediente", "Usuario", "Título", "Fecha", "Estatus", "Próxima acción", "Fecha límite", "Hora límite", "Observaciones", "Archivo", "Notificar", "Creado en"])
    for m in Movimiento.query.order_by(Movimiento.id).all():
        ws.append([
            m.id,
            m.expediente.numero if m.expediente else "",
            m.usuario.nombre if m.usuario else "",
            m.titulo,
            m.fecha,
            m.estatus,
            m.proxima_accion,
            m.fecha_limite,
            m.hora_limite,
            m.observaciones,
            m.archivo_url,
            "Sí" if m.notificar else "No",
            m.creado_en.strftime("%d/%m/%Y %H:%M") if m.creado_en else "",
        ])

    # AVISOS INTERNOS
    ws = wb.create_sheet("Avisos internos")
    ws.append(["ID", "Usuario", "Título", "Mensaje", "Enlace", "Leída", "Creado en"])
    for n in NotificacionInterna.query.order_by(NotificacionInterna.id).all():
        ws.append([
            n.id,
            n.usuario.nombre if n.usuario else "",
            n.titulo,
            n.mensaje,
            n.enlace,
            "Sí" if n.leida else "No",
            n.creado_en.strftime("%d/%m/%Y %H:%M") if n.creado_en else "",
        ])

    # AVISOS DE EXPEDIENTE
    ws = wb.create_sheet("Avisos expediente")
    ws.append(["ID", "Expediente", "Usuario", "Título", "Mensaje", "Fecha aviso", "Creado en"])
    for av in AvisoExpediente.query.order_by(AvisoExpediente.id).all():
        ws.append([
            av.id,
            av.expediente.numero if av.expediente else "",
            av.usuario.nombre if av.usuario else "",
            av.titulo,
            av.mensaje,
            av.fecha_aviso,
            av.creado_en.strftime("%d/%m/%Y %H:%M") if av.creado_en else "",
        ])

    # CITAS / PROSPECTOS
    ws = wb.create_sheet("Citas")
    ws.append(["ID", "Cliente", "Nombre", "Teléfono", "Correo", "Materia", "Tipo asunto", "Fecha", "Hora", "Documentación solicitada", "Documentación original", "Estatus", "Observaciones", "Creado en"])
    for ci in CitaProspecto.query.order_by(CitaProspecto.id).all():
        ws.append([ci.id, ci.cliente.nombre if ci.cliente else "", ci.nombre, ci.telefono, ci.correo, ci.materia, ci.tipo_asunto, ci.fecha, ci.hora, ci.documentacion_solicitada, ci.documentacion_original, ci.estatus, ci.observaciones, ci.creado_en.strftime("%d/%m/%Y %H:%M") if ci.creado_en else ""])

    # USUARIOS
    ws = wb.create_sheet("Usuarios")
    ws.append(["ID", "Nombre", "Usuario", "Correo", "Rol", "Área", "Foto", "Activo", "Creado en"])
    for u in Usuario.query.order_by(Usuario.id).all():
        ws.append([
            u.id,
            u.nombre,
            u.usuario,
            u.correo,
            u.rol,
            u.area,
            u.foto_url,
            "Sí" if u.activo else "No",
            u.creado_en.strftime("%d/%m/%Y %H:%M") if u.creado_en else "",
        ])

    # EXPEDIENTES COMPARTIDOS
    ws = wb.create_sheet("Compartidos")
    ws.append(["ID", "Expediente", "Usuario con acceso", "Permiso", "Creado en"])
    for comp in Compartido.query.order_by(Compartido.id).all():
        ws.append([
            comp.id,
            comp.expediente.numero if comp.expediente else "",
            comp.usuario.nombre if comp.usuario else "",
            comp.permiso,
            comp.creado_en.strftime("%d/%m/%Y %H:%M") if comp.creado_en else "",
        ])

    # BITÁCORA
    ws = wb.create_sheet("Bitacora")
    ws.append(["ID", "Expediente", "Usuario", "Acción", "Detalle", "Creado en"])
    for b in Bitacora.query.order_by(Bitacora.id).all():
        ws.append([
            b.id,
            b.expediente.numero if b.expediente else "",
            b.usuario.nombre if b.usuario else "",
            b.accion,
            b.detalle,
            b.creado_en.strftime("%d/%m/%Y %H:%M") if b.creado_en else "",
        ])

    # Ajuste visual básico de columnas
    for hoja in wb.worksheets:
        for columna in hoja.columns:
            letra = columna[0].column_letter
            maximo = 12
            for celda in columna:
                if celda.value is not None:
                    maximo = max(maximo, min(len(str(celda.value)) + 2, 45))
            hoja.column_dimensions[letra].width = maximo

    archivo = BytesIO()
    wb.save(archivo)
    archivo.seek(0)

    nombre = "ALESI_RESPALDO_" + local_now().strftime("%Y%m%d_%H%M") + ".xlsx"

    return send_file(
        archivo,
        as_attachment=True,
        download_name=nombre,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@app.route("/health")
def health():
    return "OK", 200


init_db()


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "notificaciones":
        print("Notificaciones procesadas:", revisar_y_enviar_notificaciones())
    else:
        app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)
